"""Phase-10 smoke for Pavement Condition Survey + Distress Assessment.

Headless. Asserts:
  1. PCI engine sanity:
       - empty input -> PCI = 100 ("Excellent")
       - severity escalation strictly increases the deduct
       - extent escalation strictly increases the deduct
       - PCI is floored at 0
  2. Rehab-recommendation lookup returns a placeholder treatment + IRC ref
     for every (type, severity) cell, and tags is_placeholder=True.
  3. DB round-trip (save / latest / cascade delete with Project).
  4. Standalone condition_docx is built and contains the PLACEHOLDER
     banner, ASTM-D6433 / IRC:82-1982 source tags and distress rows.
  5. Combined report includes the Pavement Condition Survey section.
  6. Panel end-to-end: ConditionSurveyPanel.compute / save / export.
  7. Reserved-fields banner present (image / AI / GIS placeholders).
"""
from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
_tmp = Path(tempfile.mkdtemp())
_db_path = _tmp / "phase10.db"

import app.config as _cfg  # noqa: E402
_cfg.DB_PATH = _db_path

from docx import Document  # noqa: E402
from PySide6.QtWidgets import QApplication, QFileDialog, QMessageBox  # noqa: E402
QMessageBox.information = staticmethod(lambda *a, **k: 0)
QMessageBox.warning     = staticmethod(lambda *a, **k: 0)
QMessageBox.critical    = staticmethod(lambda *a, **k: 0)

from app.core import (  # noqa: E402
    DISTRESS_TYPES,
    SEVERITY_LEVELS,
    ConditionSurveyInput,
    DistressRecord,
    compute_condition_survey,
    recommend_rehab,
)
from app.db.repository import Database  # noqa: E402
from app.reports import (  # noqa: E402
    CombinedReportContext,
    ConditionReportContext,
    build_combined_report,
    build_condition_docx,
)


def _txt(p: Path) -> str:
    doc = Document(str(p))
    out = [pp.text for pp in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                out.append(c.text)
    return "\n".join(out)


def main() -> int:
    print("=== PCI engine sanity ===")
    # Empty input -> PCI = 100
    r = compute_condition_survey(ConditionSurveyInput(records=()))
    assert r.pci_score == 100.0, r.pci_score
    assert r.condition_category == "Excellent"
    assert r.is_placeholder is True, "Phase 10 always sets is_placeholder=True"
    print(f"  [PASS] empty survey -> PCI=100 cat={r.condition_category!r}")

    # Severity escalation
    base_inputs = lambda sev: ConditionSurveyInput(  # noqa: E731
        records=(DistressRecord(distress_type="cracking", severity=sev,
                                length_m=200.0),)
    )
    pci_low    = compute_condition_survey(base_inputs("low")).pci_score
    pci_med    = compute_condition_survey(base_inputs("medium")).pci_score
    pci_high   = compute_condition_survey(base_inputs("high")).pci_score
    assert pci_low > pci_med > pci_high, (pci_low, pci_med, pci_high)
    print(f"  [PASS] severity climbs deduct: low={pci_low:.1f} med={pci_med:.1f} high={pci_high:.1f}")

    # Extent escalation
    pci_short = compute_condition_survey(ConditionSurveyInput(
        records=(DistressRecord(distress_type="rutting", severity="medium",
                                area_m2=5.0),))).pci_score
    pci_long  = compute_condition_survey(ConditionSurveyInput(
        records=(DistressRecord(distress_type="rutting", severity="medium",
                                area_m2=500.0),))).pci_score
    assert pci_short > pci_long, (pci_short, pci_long)
    print(f"  [PASS] extent climbs deduct: 5 m^2={pci_short:.1f} -> 500 m^2={pci_long:.1f}")

    # Floor at 0
    pci_floor = compute_condition_survey(ConditionSurveyInput(records=(
        DistressRecord(distress_type="potholes", severity="high", count=500),
    ))).pci_score
    assert pci_floor == 0.0, pci_floor
    print(f"  [PASS] PCI is floored at 0 under extreme distress")

    print("\n=== Rehab recommendations ===")
    for code in DISTRESS_TYPES:
        for sev in SEVERITY_LEVELS:
            rec = recommend_rehab(code, sev)
            assert rec.treatment, f"empty treatment for {code}/{sev}"
            assert rec.is_placeholder is True
            assert rec.reference.code_id, f"missing source for {code}/{sev}"
    print(f"  [PASS] {len(DISTRESS_TYPES) * len(SEVERITY_LEVELS)} (type, severity) "
          f"cells return placeholder treatment + IRC ref")

    print("\n=== DB roundtrip + cascade ===")
    db = Database(_db_path)
    proj = db.create_project(work_name="Phase 10 Smoke")
    result = compute_condition_survey(ConditionSurveyInput(
        work_name=proj.work_name,
        surveyed_by="QA",
        survey_date="2026-05-15",
        lane_id="LHS",
        chainage_from_km=12.0, chainage_to_km=12.5,
        records=(
            DistressRecord("cracking", "medium", length_m=120.0),
            DistressRecord("potholes", "high", count=8),
            DistressRecord("ravelling", "low", area_m2=15.0),
        ),
    ))
    db.save_condition_survey(project_id=proj.id, result=result)
    got = db.latest_condition_survey(proj.id)
    assert got is not None and abs(got.pci_score - result.pci_score) < 1e-6
    assert got.condition_category == result.condition_category
    print(f"  [PASS] id={got.id} PCI={got.pci_score:.2f} cat={got.condition_category!r}")

    print("\n=== Standalone Word ===")
    tpath = _tmp / "condition.docx"
    build_condition_docx(tpath, ConditionReportContext(
        work_name=proj.work_name,
        project_title=proj.work_name,
    ), result)
    txt = _txt(tpath)
    for cite in ("ASTM-D6433", "IRC:82-1982"):
        assert cite in txt, f"missing {cite}"
    assert "PAVEMENT CONDITION SURVEY" in txt
    assert "PLACEHOLDER" in txt, "PLACEHOLDER banner must appear in report"
    assert "Reserved for Future Expansion" in txt
    assert "Potholes" in txt and "Cracking" in txt
    print(f"  [PASS] {tpath.name} ({tpath.stat().st_size} bytes)")

    print("\n=== Combined report includes Condition ===")
    cpath = _tmp / "combined.docx"
    out, included = build_combined_report(
        cpath, db, proj.id,
        CombinedReportContext(project_title=proj.work_name,
                              work_name=proj.work_name),
    )
    assert "Pavement Condition Survey" in included, included
    combined_txt = _txt(out)
    assert "PAVEMENT CONDITION SURVEY" in combined_txt
    assert "ASTM-D6433" in combined_txt
    print(f"  [PASS] sections: {included}")

    print("\n=== Panel end-to-end ===")
    app = QApplication.instance() or QApplication(sys.argv)
    panel_target = _tmp / "PanelCondition.docx"
    QFileDialog.getSaveFileName = staticmethod(lambda *a, **k: (str(panel_target), "Word"))
    from app.ui.main_window import MainWindow
    w = MainWindow()
    proj2 = w.db.create_project(work_name="Panel Condition")
    w._current_project_id = proj2.id
    w.condition.set_project(proj2.id, proj2.work_name)
    # Add two distress rows
    w.condition._add_row(distress_code="cracking", severity="high",
                         length_m=150.0)
    w.condition._add_row(distress_code="potholes", severity="medium",
                         count=4)
    w.condition._on_compute()
    last = w.condition.last_result()
    assert last is not None, "panel compute produced no result"
    assert last.pci_score < 100.0
    w.condition._on_save()
    w.condition.btn_export.click()
    assert panel_target.exists(), "panel export did not write file"
    print(f"  [PASS] panel export wrote {panel_target.stat().st_size} bytes  "
          f"(PCI={last.pci_score:.2f})")

    print("\n=== Cascade delete ===")
    db.delete_project(proj.id)
    assert db.latest_condition_survey(proj.id) is None
    print("  [PASS] condition-survey row removed with project")

    # ---------------------------------------------------------------------
    # Stabilization-sprint additions (PCI calibration / F3 / F5 / F6 /
    # reserved-field round-trip).
    # ---------------------------------------------------------------------
    print("\n=== Centralized PCI calibration ===")
    from app.core import (
        PCICalibration,
        get_pci_calibration,
        set_pci_calibration,
        reset_pci_calibration,
    )
    sample = ConditionSurveyInput(records=(
        DistressRecord("cracking", "high", length_m=200.0),
    ))
    pci_default = compute_condition_survey(sample).pci_score

    # Swap to a calibration that doubles all severity weights -> larger deduct
    cur = get_pci_calibration()
    hot = PCICalibration(
        label="smoke-test-doubled-severities",
        severity_weights={k: 2.0 * v for k, v in cur.severity_weights.items()},
        distress_weights=dict(cur.distress_weights),
        extent_divisor_length_m=cur.extent_divisor_length_m,
        extent_divisor_area_m2=cur.extent_divisor_area_m2,
        extent_divisor_count=cur.extent_divisor_count,
        is_placeholder=True,
    )
    set_pci_calibration(hot)
    pci_doubled = compute_condition_survey(sample).pci_score
    assert pci_doubled < pci_default, (pci_default, pci_doubled)
    reset_pci_calibration()
    pci_reset = compute_condition_survey(sample).pci_score
    assert pci_reset == pci_default, (pci_default, pci_reset)
    print(f"  [PASS] calibration swap: default PCI={pci_default:.2f} -> "
          f"doubled PCI={pci_doubled:.2f}, reset PCI={pci_reset:.2f}")

    print("\n=== F5 — spec-driven target air voids ===")
    from app.core import MIX_SPECS, spec_target_air_voids
    # DBM-II spec is 3-5 -> midpoint = 4.0 (matches legacy constant)
    av_dbm2 = spec_target_air_voids(MIX_SPECS["DBM-II"])
    assert abs(av_dbm2 - 4.0) < 1e-9, av_dbm2
    av_none = spec_target_air_voids(None)
    assert av_none == 4.0
    print(f"  [PASS] DBM-II midpoint={av_dbm2:g}, fallback={av_none:g}")

    print("\n=== F3 — compaction blows visible in Word report ===")
    # Build a fresh DBM-II mix-design and assert the new project-info row
    # carries the 75 blows/face line; the OBC section cites the spec range.
    from app.core import (
        GradationInput, MixDesignInput, compute_mix_design,
    )
    from app.core.models import ProjectInfo
    # Reuse the InputsPanel demo defaults to compose a coherent dataset.
    from app.ui.widgets.inputs_panel import InputsPanel
    ipanel = InputsPanel()
    ipanel.set_mix_type("DBM-II")
    payload = ipanel.collect_all()
    grad = payload["gradation"]
    grad_no_cement = GradationInput(
        sieve_sizes_mm=grad.sieve_sizes_mm,
        pass_pct=grad.pass_pct,
        blend_ratios={k: v for k, v in grad.blend_ratios.items() if k.lower() != "cement"},
        spec_lower=grad.spec_lower,
        spec_upper=grad.spec_upper,
    )
    coarse, fine, bit = payload["spgr"]
    gmm_in = payload["gmm_tab"].collect(bitumen_sg=0.0)
    md_in = MixDesignInput(
        project=ProjectInfo(mix_type="DBM-II", work_name="phase9-followup", client=""),
        gradation=grad_no_cement,
        sg_coarse=coarse, sg_fine=fine, sg_bitumen=bit,
        gmb=payload["gmb"], gmm=gmm_in,
        stability_flow=payload["stability_flow"],
    )
    md_res = compute_mix_design(md_in)
    # Target air voids should be the spec midpoint (3-5 -> 4.0).
    assert abs(md_res.obc.target_air_voids_pct - 4.0) < 1e-6, \
        md_res.obc.target_air_voids_pct
    print(f"  [PASS] OBCResult.target_air_voids_pct={md_res.obc.target_air_voids_pct:g} "
          f"(spec midpoint, not flat constant)")

    from app.reports.word_report import build_mix_design_docx, ReportContext
    from app.graphs import build_chart_set
    md_doc = _tmp / "mix_phase9_followup.docx"
    build_mix_design_docx(
        md_doc,
        ReportContext(project_title="phase9-followup", mix_type_key="DBM-II"),
        md_res,
        build_chart_set(md_res.summary, md_res.obc),
        material_calc=None,
    )
    md_txt = _txt(md_doc)
    assert "Compaction (each face)" in md_txt, "F3 project-info row missing"
    assert "75 blows" in md_txt, "F3 compaction blows value missing"
    assert "midpoint of" in md_txt, "F5 spec-driven OBC annotation missing"
    print("  [PASS] Word report cites compaction blows + spec-driven AV midpoint")

    print("\n=== F6 — project-form placeholder suffix ===")
    from app.ui.widgets.project_form import ProjectForm
    pf = ProjectForm(w.db)
    placeholder_items = []
    verified_items = []
    for i in range(pf.mix_type.count()):
        code = pf.mix_type.itemData(i)
        if code is None:
            continue
        label = pf.mix_type.itemText(i)
        if "[placeholder]" in label:
            placeholder_items.append(code)
        else:
            verified_items.append(code)
    assert "DBM-II" in verified_items, "DBM-II must NOT carry [placeholder] suffix"
    assert "SMA" in placeholder_items, "SMA must carry [placeholder] suffix"
    assert "BC-II" in verified_items, "BC-II must NOT carry [placeholder] suffix"
    print(f"  [PASS] project-form annotates {len(placeholder_items)} placeholder mixes")

    print("\n=== Reserved-field round-trip (Phase 11/12 hook) ===")
    proj3 = w.db.create_project(work_name="Reserved-field round-trip")
    w._current_project_id = proj3.id
    w.condition.set_project(proj3.id, proj3.work_name)
    w.condition.ed_image_paths.setPlainText("photos/a.jpg\nphotos/b.jpg")
    w.condition.ed_ai_hint.setText("ravelling visible on inner wheel path")
    w.condition.ed_gis_geojson.setPlainText('{"type":"Point","coordinates":[78.0,17.5]}')
    w.condition._add_row(distress_code="ravelling", severity="medium",
                         area_m2=12.0)
    w.condition._on_compute()
    w.condition._on_save()
    # Reload the project freshly through set_project to prove round-trip.
    w.condition.set_project(None)
    w.condition.set_project(proj3.id, proj3.work_name)
    assert w.condition.ed_image_paths.toPlainText().splitlines() == [
        "photos/a.jpg", "photos/b.jpg"]
    assert w.condition.ed_ai_hint.text() == "ravelling visible on inner wheel path"
    assert "78.0" in w.condition.ed_gis_geojson.toPlainText()
    print("  [PASS] image_paths / ai_hint / gis_geojson round-trip via save+reload")

    print("\nPHASE 10 SMOKE: ALL OK")
    return 0


if __name__ == "__main__":
    sys.exit(main())
