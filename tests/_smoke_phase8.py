"""Phase-8 headless smoke for Traffic / ESAL / MSA module."""
from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
_tmp = Path(tempfile.mkdtemp())
_db_path = _tmp / "phase8.db"

import app.config as _cfg  # noqa: E402
_cfg.DB_PATH = _db_path

from docx import Document  # noqa: E402
from PySide6.QtWidgets import QApplication, QFileDialog, QMessageBox  # noqa: E402
QMessageBox.information = staticmethod(lambda *a, **k: 0)
QMessageBox.warning     = staticmethod(lambda *a, **k: 0)
QMessageBox.critical    = staticmethod(lambda *a, **k: 0)

from app.core import (  # noqa: E402
    TrafficInput, compute_traffic_analysis, traffic_category,
    vdf_preset, ldf_preset,
)
from app.db.repository import Database  # noqa: E402
from app.reports import (  # noqa: E402
    CombinedReportContext, TrafficReportContext,
    build_combined_report, build_traffic_docx,
)


def _txt(p: Path) -> str:
    doc = Document(str(p))
    out = [pp.text for pp in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                out.append(c.text)
    return "\n".join(out)


def main() -> int:
    print("=== Engine sanity ===")
    # Preset lookups
    assert vdf_preset("Plain", 2000) == 4.8
    assert vdf_preset("Hilly", 100) == 0.5
    assert ldf_preset("Two-lane carriageway") == 0.50
    # Category thresholds
    assert traffic_category(3) == "Low (< 5 MSA)"
    assert traffic_category(20) == "Moderate (5-30 MSA)" or traffic_category(20).startswith("Moderate")
    # Full compute — matches structural_design (no override)
    r = compute_traffic_analysis(TrafficInput(
        initial_cvpd=2000, growth_rate_pct=7.5,
        design_life_years=15, terrain="Plain",
        lane_config="Two-lane carriageway",
        road_category="NH / SH",
    ))
    # VDF preset for Plain >1500 = 4.8, LDF for two-lane carriageway = 0.50
    # MSA = 365 * 2000 * GF(7.5%,15) * 0.50 * 4.8 / 1e6
    # GF = ((1.075^15 - 1) / 0.075) = 26.121
    # MSA = 365 * 2000 * 26.121 * 0.50 * 4.8 / 1e6 = 45.764
    assert abs(r.design_msa - 45.764) < 0.05, r.design_msa
    assert abs(r.aashto_esal - r.design_msa * 1e6) < 1
    assert r.traffic_category.startswith("Heavy"), r.traffic_category
    print(f"  OK -- MSA={r.design_msa:.2f} ESAL={r.aashto_esal:,.0f} cat={r.traffic_category}")

    # Override path
    r2 = compute_traffic_analysis(TrafficInput(
        initial_cvpd=500, growth_rate_pct=5.0,
        design_life_years=10, vdf=3.0, ldf=0.6,
    ))
    assert r2.vdf_used == 3.0 and r2.ldf_used == 0.6
    print(f"  OK -- override VDF/LDF respected: MSA={r2.design_msa:.2f}")

    print("\n=== DB roundtrip + cascade ===")
    db = Database(_db_path)
    proj = db.create_project(work_name="Phase 8 Smoke")
    db.save_traffic_analysis(project_id=proj.id, result=r)
    got = db.latest_traffic_analysis(proj.id)
    assert got is not None and abs(got.design_msa - r.design_msa) < 1e-6
    print(f"  OK -- id={got.id} MSA={got.design_msa:.2f}")

    print("\n=== Standalone Word ===")
    tpath = _tmp / "traffic.docx"
    build_traffic_docx(tpath, TrafficReportContext(work_name=proj.work_name), r)
    txt = _txt(tpath)
    for cite in ("IRC:37-2018", "AASHTO-1993"):
        assert cite in txt, f"missing {cite}"
    assert "TRAFFIC / ESAL / MSA" in txt
    assert "Reserved for Future Expansion" in txt   # placeholders rendered
    print(f"  OK -- {tpath.name} ({tpath.stat().st_size} bytes)")

    print("\n=== Combined report includes Traffic ===")
    cpath = _tmp / "combined.docx"
    out, included = build_combined_report(
        cpath, db, proj.id,
        CombinedReportContext(project_title=proj.work_name,
                              work_name=proj.work_name),
    )
    assert "Traffic / ESAL / MSA Analysis" in included, included
    assert "TRAFFIC / ESAL / MSA" in _txt(out)
    print(f"  OK -- sections: {included}")

    print("\n=== Panel end-to-end ===")
    app = QApplication.instance() or QApplication(sys.argv)
    panel_target = _tmp / "PanelTraffic.docx"
    QFileDialog.getSaveFileName = staticmethod(lambda *a, **k: (str(panel_target), "Word"))
    from app.ui.main_window import MainWindow
    w = MainWindow()
    proj2 = w.db.create_project(work_name="Panel Traffic")
    w._current_project_id = proj2.id
    w.traffic.set_project(proj2.id, proj2.work_name)
    w.traffic._on_compute()
    assert w.traffic.last_result() is not None
    w.traffic._on_save()
    w.traffic.btn_export.click()
    assert panel_target.exists()
    print(f"  OK -- panel export wrote {panel_target.stat().st_size} bytes")

    print("\n=== Cascade delete ===")
    db.delete_project(proj.id)
    assert db.latest_traffic_analysis(proj.id) is None
    print("  OK")

    print("\nPHASE 8 SMOKE: ALL OK")
    return 0


if __name__ == "__main__":
    sys.exit(main())
