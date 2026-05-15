"""Phase-15 Priority-3 smoke — structural report modernization.

Verifies the structural Word report renders both flavours of the
mechanistic-checks section:

  * legacy fallback   — StructuralResult with no
                         ``mechanistic_validation`` reference still
                         emits the original "Check / Status" table
                         and the standard fatigue / rutting strings.
  * modernized path   — StructuralResult with a Phase-14
                         ``MechanisticValidationSummary`` attached
                         renders the rich mechanistic section
                         (verdicts, calibration table, IRC:37 cl.
                         6.4 citations).
  * refused mechanistic input renders the [REFUSED] banner with the
    Phase-14 refusal reason verbatim.

Pure-Python + python-docx; no Qt, no DB.
"""
from __future__ import annotations

import dataclasses
import tempfile
from pathlib import Path

from docx import Document

from app.core import (
    IITPavePavementLayer,
    IITPavePointResult,
    LABEL_FATIGUE,
    LABEL_RUTTING,
    MechanisticResult,
    MechanisticValidationInput,
    PavementStructure,
    StructuralInput,
    compute_mechanistic_validation,
    compute_structural_design,
)
from app.core.code_refs import CodeRef
from app.reports import (
    StructuralReportContext,
    build_structural_docx,
)


def _structure(e_bc_mpa: float = 3000.0) -> PavementStructure:
    return PavementStructure(layers=(
        IITPavePavementLayer(name="BC",  material="BC",  modulus_mpa=e_bc_mpa,
                             thickness_mm=40.0),
        IITPavePavementLayer(name="DBM", material="DBM", modulus_mpa=2500.0,
                             thickness_mm=80.0),
        IITPavePavementLayer(name="GSB", material="GSB", modulus_mpa=300.0,
                             thickness_mm=200.0),
        IITPavePavementLayer(name="Subgrade", material="Subgrade",
                             modulus_mpa=50.0, thickness_mm=None),
    ))


def _mech_result(*, is_placeholder: bool,
                 eps_t_microstrain: float = 150.0,
                 eps_v_microstrain: float = 300.0,
                 source: str = "external_exe") -> MechanisticResult:
    bt = IITPavePointResult(
        z_mm=120.0, r_mm=0.0,
        sigma_z_mpa=0.5, sigma_r_mpa=0.1, sigma_t_mpa=0.1,
        epsilon_z_microstrain=-eps_t_microstrain,
        epsilon_r_microstrain=eps_t_microstrain,
        epsilon_t_microstrain=eps_t_microstrain,
    )
    sg = IITPavePointResult(
        z_mm=320.0, r_mm=0.0,
        sigma_z_mpa=0.05, sigma_r_mpa=0.01, sigma_t_mpa=0.01,
        epsilon_z_microstrain=eps_v_microstrain,
        epsilon_r_microstrain=-eps_v_microstrain * 0.3,
        epsilon_t_microstrain=-eps_v_microstrain * 0.3,
    )
    return MechanisticResult(
        point_results=(bt, sg),
        references=(CodeRef("IRC:37-2018", "cl. 6.2",
                            "Multi-layered elastic analysis"),),
        is_placeholder=is_placeholder,
        source=source,
    )


def _docx_text(p: Path) -> str:
    doc = Document(str(p))
    out = [pp.text for pp in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                out.append(c.text)
    return "\n".join(out)


def main() -> int:
    tmp = Path(tempfile.mkdtemp())
    ctx = StructuralReportContext(
        project_title="Phase 15 P3 Smoke — Structural Modernization",
        work_name="NH-007 Test Section",
        work_order_no="WO/2026/P15P3",
        client="QA Client",
        agency="QA Agency",
        submitted_by="QA Engineer",
    )
    base = compute_structural_design(StructuralInput(
        initial_cvpd=2000.0, growth_rate_pct=5.0, design_life_years=15,
        subgrade_cbr_pct=8.0,
    ))

    # ---------------------------------------------------------------
    # 1) Legacy fallback path — mechanistic_validation = None
    # ---------------------------------------------------------------
    print("=== 1) Legacy fallback path ===")
    legacy = dataclasses.replace(base,
                                 fatigue_check="Not performed.",
                                 rutting_check="Not performed.")
    assert legacy.mechanistic_validation is None
    p_legacy = tmp / "struct_legacy.docx"
    build_structural_docx(p_legacy, ctx, legacy)
    txt_legacy = _docx_text(p_legacy)
    assert "FLEXIBLE PAVEMENT STRUCTURAL DESIGN" in txt_legacy
    # The legacy table heading rows must appear.
    assert "Bituminous-layer fatigue" in txt_legacy
    assert "Subgrade rutting" in txt_legacy
    assert "Not performed." in txt_legacy
    # Modern section markers must NOT be present.
    assert "MECHANISTIC VALIDATION (IRC:37-2018 cl. 6.4)" not in txt_legacy
    assert "IRC37_PLACEHOLDER_80pct" not in txt_legacy
    print(f"  [PASS] {p_legacy.name} ({p_legacy.stat().st_size} bytes)")

    # ---------------------------------------------------------------
    # 2) Modernized path — mechanistic_validation present (verdict).
    # ---------------------------------------------------------------
    print("\n=== 2) Modernized verdict path ===")
    verdict_summary = compute_mechanistic_validation(
        MechanisticValidationInput(
            mech_result=_mech_result(is_placeholder=False),
            structure=_structure(),
            design_msa=base.design_msa,
            point_labels=(LABEL_FATIGUE, LABEL_RUTTING),
        )
    )
    assert verdict_summary.refused is False
    modern = dataclasses.replace(base, mechanistic_validation=verdict_summary)
    p_modern = tmp / "struct_modern_verdict.docx"
    build_structural_docx(p_modern, ctx, modern)
    txt_modern = _docx_text(p_modern)
    # Rich section markers present.
    assert "Validation Summary" in txt_modern
    assert "Fatigue (Bottom of Bituminous-Bound Layer)" in txt_modern
    assert "Rutting (Top of Subgrade)" in txt_modern
    assert "IRC37_PLACEHOLDER_80pct" in txt_modern
    assert verdict_summary.fatigue.verdict in txt_modern
    assert verdict_summary.rutting.verdict in txt_modern
    # Calibration is still placeholder, so the placeholder banner fires.
    assert "[PLACEHOLDER]" in txt_modern
    # Legacy "Not performed." must NOT appear in the modernized doc.
    assert "Not performed." not in txt_modern
    print(f"  [PASS] {p_modern.name} ({p_modern.stat().st_size} bytes); "
          f"fatigue={verdict_summary.fatigue.verdict} "
          f"rutting={verdict_summary.rutting.verdict}")

    # ---------------------------------------------------------------
    # 3) Modernized path — refused (placeholder mechanistic input).
    # ---------------------------------------------------------------
    print("\n=== 3) Modernized refused path ===")
    refused_summary = compute_mechanistic_validation(
        MechanisticValidationInput(
            mech_result=_mech_result(is_placeholder=True, source="stub"),
            structure=_structure(),
            design_msa=base.design_msa,
            point_labels=(LABEL_FATIGUE, LABEL_RUTTING),
        )
    )
    assert refused_summary.refused is True
    refused_struct = dataclasses.replace(base,
                                         mechanistic_validation=refused_summary)
    p_refused = tmp / "struct_modern_refused.docx"
    build_structural_docx(p_refused, ctx, refused_struct)
    txt_refused = _docx_text(p_refused)
    assert "[REFUSED]" in txt_refused
    # Verbatim refusal reason from Phase 14.
    assert "MechanisticResult.is_placeholder=True" in txt_refused
    # Rich section still renders despite the refusal (banner + KV with —).
    assert "Validation Summary" in txt_refused
    # Legacy fallback must NOT appear.
    assert "Not performed." not in txt_refused
    print(f"  [PASS] {p_refused.name} ({p_refused.stat().st_size} bytes)")

    print("\nPHASE 15 STRUCTURAL MODERN SMOKE: ALL OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
