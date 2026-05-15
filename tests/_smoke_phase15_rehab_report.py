"""Phase-15 Priority-2 smoke — rehab recommendation Word report.

Covers:

  1. Standalone build_rehab_docx with NO recommendations
     (Excellent-PCI survey) -> placeholder banner, empty-state note,
     references block still cites IRC governing codes.
  2. Standalone build_rehab_docx with a populated synthesis
     (Fair PCI + traffic) -> prioritized table + per-recommendation
     detail blocks + IRC source-tags inline.
  3. Combined-report integration: a project with a saved condition
     survey + traffic analysis must produce a Word doc whose included
     sections list contains "Rehabilitation Recommendations" AND the
     rendered text carries the section title and recommendation labels.

Pure-Python + python-docx; uses a real Database backed by a tempdir
sqlite file so the combined-report flow exercises the full
``build_combined_report`` path.
"""
from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

_tmp = Path(tempfile.mkdtemp())

import app.config as _cfg  # noqa: E402
_cfg.DB_PATH = _tmp / "phase15_p2.db"

from docx import Document  # noqa: E402

from app.core import (  # noqa: E402
    ConditionSurveyInput,
    DistressRecord,
    RecommendationContext,
    TrafficInput,
    compute_condition_survey,
    compute_rehab_recommendations,
    compute_traffic_analysis,
)
from app.db.repository import Database  # noqa: E402
from app.reports import (  # noqa: E402
    CombinedReportContext,
    RehabReportContext,
    build_combined_report,
    build_rehab_docx,
)


def _docx_text(p: Path) -> str:
    doc = Document(str(p))
    out = [pp.text for pp in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                out.append(c.text)
    return "\n".join(out)


def main() -> int:
    ctx = RehabReportContext(
        project_title="Phase 15 P2 Smoke — Rehab Recommendations",
        work_name="SH-21 Strengthening (Smoke Project)",
        work_order_no="WO/2026/P15P2",
        client="QA Client",
        agency="QA Agency",
        submitted_by="QA Engineer",
    )

    # ---------------------------------------------------------------
    # 1) Empty / Excellent survey  -> ROUTINE_MAINTENANCE only +
    #    placeholder banner + references still present.
    # ---------------------------------------------------------------
    print("=== 1) Excellent survey -> routine recommendation ===")
    empty_cond = compute_condition_survey(ConditionSurveyInput())
    s_empty = compute_rehab_recommendations(RecommendationContext(condition=empty_cond))
    p_empty = _tmp / "rehab_empty.docx"
    build_rehab_docx(p_empty, ctx, s_empty)
    txt_empty = _docx_text(p_empty)
    assert "REHABILITATION RECOMMENDATIONS" in txt_empty
    assert "[PLACEHOLDER]" in txt_empty
    # ROUTINE_MAINTENANCE label rendered.
    assert "Routine maintenance" in txt_empty
    # References block cites IRC governing codes.
    assert "References" in txt_empty
    for cite in ("IRC:82-1982", "IRC:81-1997", "IRC:115"):
        assert cite in txt_empty, f"missing {cite} in empty-survey report"
    print(f"  [PASS] {p_empty.name} ({p_empty.stat().st_size} bytes)")

    # ---------------------------------------------------------------
    # 2) Fair PCI + mid-range MSA -> Crack Sealing + Surface Treatment
    #    + Micro Surfacing among the recommendations.
    # ---------------------------------------------------------------
    print("\n=== 2) Fair PCI + mid-range MSA -> multi-recommendation ===")
    cond_fair = compute_condition_survey(ConditionSurveyInput(
        work_name="Fair PCI fixture",
        records=(
            DistressRecord("rutting", "medium", area_m2=80.0),
            DistressRecord("cracking", "low", length_m=80.0),
            DistressRecord("ravelling", "low", area_m2=20.0),
        ),
    ))
    traffic_mid = compute_traffic_analysis(TrafficInput(
        initial_cvpd=800.0, growth_rate_pct=5.0, design_life_years=15,
        terrain="Plain", lane_config="Two-lane carriageway",
    ))
    s_fair = compute_rehab_recommendations(RecommendationContext(
        condition=cond_fair, traffic=traffic_mid,
    ))
    p_fair = _tmp / "rehab_fair.docx"
    build_rehab_docx(p_fair, ctx, s_fair)
    txt_fair = _docx_text(p_fair)
    assert "Prioritized Treatment Recommendations" in txt_fair
    assert "Recommendation Detail" in txt_fair
    # At least one expected category label.
    assert "Micro-surfacing" in txt_fair or "Slurry seal" in txt_fair, txt_fair[:500]
    assert "Crack sealing" in txt_fair, "low-severity cracking should trigger Crack sealing"
    # Source-tags rendered inline (any one of these IRC ids).
    assert any(t in txt_fair for t in ("IRC:82-1982", "IRC:SP:101", "IRC:SP:81"))
    print(f"  [PASS] {p_fair.name} ({p_fair.stat().st_size} bytes); "
          f"recs={len(s_fair.recommendations)}")

    # ---------------------------------------------------------------
    # 3) Combined-report integration via build_combined_report.
    # ---------------------------------------------------------------
    print("\n=== 3) Combined-report integration ===")
    db = Database(_cfg.DB_PATH)
    proj = db.create_project(work_name="Phase 15 P2 combined", mix_type="DBM-II")
    db.save_condition_survey(project_id=proj.id, result=cond_fair)
    db.save_traffic_analysis(project_id=proj.id,
                             result=traffic_mid)
    combined_path = _tmp / "phase15_p2_combined.docx"
    out_path, included = build_combined_report(
        combined_path,
        db,
        proj.id,
        CombinedReportContext(
            project_title=proj.work_name,
            work_name=proj.work_name,
            client="QA Client",
            agency="QA Agency",
            submitted_by="QA Engineer",
        ),
    )
    assert "Rehabilitation Recommendations" in included, included
    txt_combined = _docx_text(out_path)
    assert "REHABILITATION RECOMMENDATIONS" in txt_combined
    assert "Prioritized Treatment Recommendations" in txt_combined
    # ToC row present.
    assert "Rehabilitation Recommendations" in txt_combined
    assert "IRC:82-1982" in txt_combined
    # Combined report must still keep its pre-existing sections.
    assert "TRAFFIC / ESAL / MSA ANALYSIS" in txt_combined
    assert "PAVEMENT CONDITION SURVEY" in txt_combined
    print(f"  [PASS] {out_path.name} ({out_path.stat().st_size} bytes); "
          f"sections={included}")

    print("\nPHASE 15 REHAB REPORT SMOKE: ALL OK")
    return 0


if __name__ == "__main__":
    sys.exit(main())
