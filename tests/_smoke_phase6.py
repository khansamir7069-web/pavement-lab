"""Phase-6 headless smoke for the module-aware report builders.

Exercises:
  * Engine sanity (IRC:81 lookup at tabulated MSAs).
  * Standalone Structural Word docx (IRC:37-2018).
  * Standalone Maintenance Word docx (IRC:81 / IRC:SP:100 / IRC:SP:81).
  * Combined Word docx via the DB (no live mix-design supplied).
  * Verifies that key IRC citations appear in the generated documents.

Run with:
    python -m tests._smoke_phase6
"""
from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

_tmp = Path(tempfile.mkdtemp())
_db_path = _tmp / "phase6.db"

import app.config as _cfg  # noqa: E402
_cfg.DB_PATH = _db_path

from docx import Document  # noqa: E402

from app.core import (  # noqa: E402
    ColdMixInput, MicroSurfacingInput, OverlayInput, StructuralInput,
    allowable_deflection, compute_cold_mix, compute_micro_surfacing,
    compute_overlay, compute_structural_design,
)
from app.db.repository import Database  # noqa: E402
from app.reports import (  # noqa: E402
    CombinedReportContext, MaintenanceReportContext, StructuralReportContext,
    build_combined_report, build_maintenance_docx, build_structural_docx,
)


def _all_text(path: Path) -> str:
    """Concatenate all paragraph + table cell text in a docx file."""
    doc = Document(str(path))
    out: list[str] = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                out.append(c.text)
    return "\n".join(out)


def section(title: str) -> None:
    print("\n=== " + title + " ===")


def main() -> int:
    section("Engine IRC:81 plate-1 sanity")
    for n, d in [(1, 1.50), (5, 1.10), (10, 0.90), (30, 0.65), (50, 0.60), (100, 0.50)]:
        got = allowable_deflection(n)
        assert abs(got - d) < 1e-9, f"{n} MSA -> {got} (want {d})"
        print(f"  N={n:>3} -> {got:.3f}")

    db = Database(_db_path)
    proj = db.create_project(work_name="Phase 6 Smoke Project",
                             agency="State PWD", submitted_by="Test Engineer")
    pid = proj.id
    print(f"\nCreated project id={pid}")

    # Save a structural design
    s_inp = StructuralInput(road_category="NH / SH", design_life_years=15,
                            initial_cvpd=2000, growth_rate_pct=7.5,
                            vdf=2.5, ldf=0.75, subgrade_cbr_pct=5.0)
    s_res = compute_structural_design(s_inp)
    db.save_structural_design(project_id=pid, result=s_res)
    print(f"Saved structural: MSA={s_res.design_msa:.2f} Mr={s_res.subgrade_mr_mpa:.1f}")

    # Save maintenance: all three sub-modules
    o_res = compute_overlay(OverlayInput(
        deflections_mm=(0.85, 1.05, 0.95, 1.20, 1.10),
        pavement_temp_c=28.0, bituminous_thickness_mm=100.0,
        season_factor=1.0, road_category="NH / SH",
        design_traffic_msa=20.0,
    ))
    db.save_maintenance_design(project_id=pid, sub_module="overlay", result=o_res)
    c_res = compute_cold_mix(ColdMixInput(mix_type="Dense-Graded"))
    db.save_maintenance_design(project_id=pid, sub_module="cold_mix", result=c_res)
    m_res = compute_micro_surfacing(MicroSurfacingInput(surfacing_type="Type II"))
    db.save_maintenance_design(project_id=pid, sub_module="micro_surfacing", result=m_res)
    print(f"Saved maintenance: overlay h={o_res.overlay_thickness_mm:.0f} mm  "
          f"cold residual={c_res.residual_binder_pct:.2f}%  "
          f"micro residual={m_res.residual_binder_pct:.2f}%")

    section("Standalone Structural docx (IRC:37-2018)")
    s_path = _tmp / "structural.docx"
    build_structural_docx(s_path,
                          StructuralReportContext(project_title=proj.work_name,
                                                  work_name=proj.work_name,
                                                  submitted_by="Test Engineer"),
                          s_res)
    txt = _all_text(s_path)
    assert "IRC:37-2018" in txt, "structural report missing IRC:37-2018 citation"
    assert "Cumulative Design Traffic" in txt
    assert "Subgrade Resilient Modulus" in txt
    assert f"{s_res.design_msa:.2f}" in txt
    print(f"  OK — {s_path.name} ({s_path.stat().st_size} bytes)")

    section("Standalone Maintenance docx (IRC:81 / IRC:SP:100 / IRC:SP:81)")
    m_path = _tmp / "maintenance.docx"
    build_maintenance_docx(
        m_path,
        MaintenanceReportContext(project_title=proj.work_name,
                                 work_name=proj.work_name),
        overlay=o_res, cold_mix=c_res, micro_surfacing=m_res,
    )
    txt = _all_text(m_path)
    for cite in ("IRC:81-1997", "IRC:SP:100-2014", "IRC:SP:81"):
        assert cite in txt, f"maintenance report missing citation {cite}"
    assert "550" in txt, "missing the IRC:81 overlay coefficient 550"
    assert "0.0065" in txt, "missing the IRC:81 T-correction factor"
    print(f"  OK — {m_path.name} ({m_path.stat().st_size} bytes)")

    section("Combined docx (module-aware, no live mix-design)")
    c_path = _tmp / "combined.docx"
    out_path, included = build_combined_report(
        c_path, db, pid,
        CombinedReportContext(project_title=proj.work_name,
                              work_name=proj.work_name,
                              submitted_by="Test Engineer"),
    )
    txt = _all_text(out_path)
    for cite in ("IRC:37-2018", "IRC:81-1997", "IRC:SP:100-2014", "IRC:SP:81"):
        assert cite in txt, f"combined report missing {cite}"
    assert "Contents of this Report" in txt
    assert len(included) == 4, f"expected 4 sections, got {included}"
    print(f"  OK — {out_path.name} ({out_path.stat().st_size} bytes)")
    print(f"     Sections: {included}")

    section("Combined docx — empty project should be rejected")
    empty = db.create_project(work_name="Empty project")
    try:
        build_combined_report(
            _tmp / "empty.docx", db, empty.id,
            CombinedReportContext(project_title=empty.work_name,
                                  work_name=empty.work_name),
        )
    except ValueError as e:
        print(f"  OK — refused: {e}")
    else:
        raise AssertionError("expected ValueError for project with no modules")

    section("Cascade delete: maintenance + structural rows gone")
    db.delete_project(pid)
    assert db.latest_structural_design(pid) is None
    assert db.latest_maintenance_design(pid, "overlay") is None
    print("  OK")

    print("\nPHASE 6 SMOKE: ALL OK")
    return 0


if __name__ == "__main__":
    sys.exit(main())
