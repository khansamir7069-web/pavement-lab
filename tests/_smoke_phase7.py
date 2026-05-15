"""Phase-7 headless smoke for Material Quantity Calculator.

Exercises:
  * Engine math sanity (DBM tonnage, prime/tack coats, granular).
  * Per-layer CodeRef tagging.
  * DB roundtrip + cascade-delete.
  * Standalone material-quantity Word docx.
  * Combined report now includes the BOQ section.
  * End-to-end via the panel (set_project -> compute -> save -> export).
"""
from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
_tmp = Path(tempfile.mkdtemp())
_db_path = _tmp / "phase7.db"

import app.config as _cfg  # noqa: E402
_cfg.DB_PATH = _db_path

from docx import Document  # noqa: E402
from PySide6.QtWidgets import QApplication, QFileDialog, QMessageBox  # noqa: E402
QMessageBox.information = staticmethod(lambda *a, **k: 0)
QMessageBox.warning     = staticmethod(lambda *a, **k: 0)
QMessageBox.critical    = staticmethod(lambda *a, **k: 0)

from app.core import (  # noqa: E402
    LayerInput, MaterialQuantityInput, compute_material_quantity,
    StructuralInput, compute_structural_design,
)
from app.db.repository import Database  # noqa: E402
from app.reports import (  # noqa: E402
    CombinedReportContext, MaterialQuantityReportContext,
    build_combined_report, build_material_quantity_docx,
)


def _all_text(path: Path) -> str:
    doc = Document(str(path))
    out = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                out.append(c.text)
    return "\n".join(out)


def main() -> int:
    print("=== Engine sanity ===")
    inp = MaterialQuantityInput(layers=(
        LayerInput("Prime Coat", length_m=1000, width_m=7.0),
        LayerInput("DBM",        length_m=1000, width_m=7.0, thickness_mm=75),
        LayerInput("Tack Coat",  length_m=1000, width_m=7.0),
        LayerInput("BC",         length_m=1000, width_m=7.0, thickness_mm=40),
        LayerInput("WMM",        length_m=1000, width_m=7.0, thickness_mm=250),
        LayerInput("GSB",        length_m=1000, width_m=7.0, thickness_mm=200),
    ))
    r = compute_material_quantity(inp)
    # Expected DBM tonnage: 1000*7 * 0.075 * 2.40 * 1.02 = 1285.2 t
    dbm = next(L for L in r.layers if L.inputs.layer_type == "DBM")
    assert abs(dbm.layer_tonnage_t - 1285.2) < 0.1, dbm.layer_tonnage_t
    # Expected DBM binder: 1285.2 * 4.5% = 57.834 t
    assert abs(dbm.binder_tonnage_t - 57.834) < 0.1, dbm.binder_tonnage_t
    # Prime coat: 7000 m² * 0.75 kg/m² / 1000 = 5.25 t
    pc = next(L for L in r.layers if L.inputs.layer_type == "Prime Coat")
    assert abs(pc.binder_tonnage_t - 5.25) < 1e-6, pc.binder_tonnage_t
    assert pc.layer_tonnage_t == 0.0
    # Granular: WMM 7000*0.25*2.20*1.02 = 3927 t (with waste 2%)
    wmm = next(L for L in r.layers if L.inputs.layer_type == "WMM")
    assert abs(wmm.layer_tonnage_t - 3927.0) < 0.5
    assert wmm.binder_tonnage_t == 0.0
    # All layers carry CodeRef tags
    for L in r.layers:
        assert L.code_refs, f"no refs for {L.inputs.layer_type}"
    print(f"  OK — Sumlayer={r.total_layer_tonnage_t:.1f} t  "
          f"Sumbinder={r.total_binder_tonnage_t:.1f} t")

    print("\n=== DB roundtrip + cascade delete ===")
    db = Database(_db_path)
    proj = db.create_project(work_name="Phase 7 Smoke")
    pid = proj.id
    db.save_material_quantity(project_id=pid, result=r)
    got = db.latest_material_quantity(pid)
    assert got is not None and got.total_layer_tonnage_t > 0
    print(f"  OK — saved id={got.id} totals={got.total_layer_tonnage_t:.1f}/{got.total_binder_tonnage_t:.1f} t")

    print("\n=== Standalone Word docx ===")
    mq_path = _tmp / "material_qty.docx"
    build_material_quantity_docx(
        mq_path,
        MaterialQuantityReportContext(project_title=proj.work_name,
                                      work_name=proj.work_name),
        r,
    )
    txt = _all_text(mq_path)
    for cite in ("MoRTH-500", "MoRTH-400", "IRC:111"):
        assert cite in txt, f"missing citation {cite}"
    assert "BILL OF MATERIAL QUANTITIES" in txt
    print(f"  OK — {mq_path.name} ({mq_path.stat().st_size} bytes)")

    print("\n=== Combined report includes BOQ ===")
    sd = compute_structural_design(StructuralInput())
    db.save_structural_design(project_id=pid, result=sd)
    c_path = _tmp / "combined.docx"
    out, included = build_combined_report(
        c_path, db, pid,
        CombinedReportContext(project_title=proj.work_name,
                              work_name=proj.work_name),
    )
    assert "Bill of Material Quantities" in included, included
    ctxt = _all_text(out)
    assert "BILL OF MATERIAL QUANTITIES" in ctxt
    print(f"  OK — sections: {included}")

    print("\n=== Panel end-to-end ===")
    app = QApplication.instance() or QApplication(sys.argv)
    panel_target = _tmp / "PanelMaterialQty.docx"
    QFileDialog.getSaveFileName = staticmethod(
        lambda *a, **k: (str(panel_target), "Word"))
    from app.ui.main_window import MainWindow
    w = MainWindow()
    proj2 = w.db.create_project(work_name="Panel Smoke")
    w._current_project_id = proj2.id
    w.material_qty.set_project(proj2.id, proj2.work_name)
    w.material_qty._on_compute()
    assert w.material_qty.last_result() is not None
    w.material_qty._on_save()
    assert w.db.latest_material_quantity(proj2.id) is not None
    w.material_qty.btn_export.click()
    assert panel_target.exists()
    print(f"  OK — panel-driven export wrote {panel_target.stat().st_size} bytes")

    print("\n=== Cascade delete ===")
    db.delete_project(pid)
    assert db.latest_material_quantity(pid) is None
    print("  OK")

    print("\nPHASE 7 SMOKE: ALL OK")
    return 0


if __name__ == "__main__":
    sys.exit(main())
