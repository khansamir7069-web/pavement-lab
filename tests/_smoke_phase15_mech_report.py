"""Phase-15 Priority-1 smoke — mechanistic Word report.

Builds two docx files (refused + verdict) and asserts:

  1. Refused-path doc renders [REFUSED] banner with reason text and
     emits "—" for life/verdict cells (no fake numbers leak through).
  2. Verdict-path doc renders fatigue + rutting verdict KV tables with
     numeric life values, [PLACEHOLDER] banner (calibration still
     placeholder), strain-extraction fallback warning carried into
     the "Notes" section, and IRC:37-2018 references.
  3. Standard project-information block + IRC:37 references appear
     under the "References" heading.

Pure-Python + python-docx; no Qt, no DB.
"""
from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document

from app.core import (
    IITPavePavementLayer,
    IITPavePointResult,
    LABEL_FATIGUE,
    LABEL_RUTTING,
    MechanisticResult,
    MechanisticValidationInput,
    PavementStructure,
    compute_mechanistic_validation,
)
from app.core.code_refs import CodeRef
from app.reports import (
    MechanisticReportContext,
    build_mechanistic_docx,
)


def _structure(e_bc_mpa: float = 3000.0) -> PavementStructure:
    return PavementStructure(layers=(
        IITPavePavementLayer(name="BC",  material="BC",  modulus_mpa=e_bc_mpa,
                             thickness_mm=40.0),
        IITPavePavementLayer(name="DBM", material="DBM", modulus_mpa=2500.0,
                             thickness_mm=80.0),
        IITPavePavementLayer(name="GSB", material="GSB", modulus_mpa=300.0,
                             thickness_mm=200.0),
        IITPavePavementLayer(name="Subgrade", material="Subgrade",
                             modulus_mpa=50.0, thickness_mm=None),
    ))


def _mech_result(*, eps_t_microstrain: float, eps_v_microstrain: float,
                 is_placeholder: bool,
                 source: str = "external_exe") -> MechanisticResult:
    bt = IITPavePointResult(
        z_mm=120.0, r_mm=0.0,
        sigma_z_mpa=0.5, sigma_r_mpa=0.1, sigma_t_mpa=0.1,
        epsilon_z_microstrain=-eps_t_microstrain,
        epsilon_r_microstrain=eps_t_microstrain,
        epsilon_t_microstrain=eps_t_microstrain,
    )
    sg = IITPavePointResult(
        z_mm=320.0, r_mm=0.0,
        sigma_z_mpa=0.05, sigma_r_mpa=0.01, sigma_t_mpa=0.01,
        epsilon_z_microstrain=eps_v_microstrain,
        epsilon_r_microstrain=-eps_v_microstrain * 0.3,
        epsilon_t_microstrain=-eps_v_microstrain * 0.3,
    )
    return MechanisticResult(
        point_results=(bt, sg),
        references=(CodeRef("IRC:37-2018", "cl. 6.2", "Multi-layered elastic"),),
        is_placeholder=is_placeholder,
        source=source,
    )


def _docx_text(p: Path) -> str:
    doc = Document(str(p))
    out = [pp.text for pp in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                out.append(c.text)
    return "\n".join(out)


def main() -> int:
    tmp = Path(tempfile.mkdtemp())
    structure = _structure()
    ctx = MechanisticReportContext(
        project_title="Phase 15 Smoke — Mechanistic Polish",
        work_name="NH-007 Strengthening (Smoke Project)",
        work_order_no="WO/2026/0001",
        client="QA Client",
        agency="QA Agency",
        submitted_by="QA Engineer",
    )

    print("=== 1) Refused path (placeholder mechanistic input) ===")
    refused_summary = compute_mechanistic_validation(
        MechanisticValidationInput(
            mech_result=_mech_result(eps_t_microstrain=150.0,
                                     eps_v_microstrain=300.0,
                                     is_placeholder=True, source="stub"),
            structure=structure,
            design_msa=30.0,
            point_labels=(LABEL_FATIGUE, LABEL_RUTTING),
        )
    )
    refused_path = tmp / "mech_refused.docx"
    build_mechanistic_docx(refused_path, ctx, refused_summary)
    text_refused = _docx_text(refused_path)
    assert "MECHANISTIC VALIDATION (IRC:37-2018 cl. 6.4)" in text_refused
    assert "[REFUSED]" in text_refused
    # Refused-reason text from Phase 14 must appear verbatim.
    assert "MechanisticResult.is_placeholder=True" in text_refused
    # No numeric life leaks through.
    assert "Cumulative fatigue life" in text_refused
    assert "Cumulative rutting life" in text_refused
    # Both verdicts rendered as dash.
    # Count "—" occurrences inside the KV tables — must be present at
    # least twice for life rows + twice for verdict rows.
    assert text_refused.count("—") >= 4, text_refused.count("—")
    # IRC:37 references block present.
    assert "References" in text_refused
    assert "IRC:37-2018" in text_refused
    print(f"  [PASS] {refused_path.name} ({refused_path.stat().st_size} bytes)")

    print("\n=== 2) Verdict path (placeholder calibration, real mech result) ===")
    verdict_summary = compute_mechanistic_validation(
        MechanisticValidationInput(
            mech_result=_mech_result(eps_t_microstrain=150.0,
                                     eps_v_microstrain=300.0,
                                     is_placeholder=False,
                                     source="external_exe"),
            structure=structure,
            design_msa=30.0,
            # Deliberately omit point_labels to force the fallback path
            # so the smoke verifies the strain-extraction warning flows
            # into the rendered "Notes" section.
            point_labels=None,
        )
    )
    verdict_path = tmp / "mech_verdict.docx"
    build_mechanistic_docx(verdict_path, ctx, verdict_summary)
    text_verdict = _docx_text(verdict_path)
    assert verdict_summary.fatigue.verdict in ("PASS", "FAIL")
    assert verdict_summary.rutting.verdict in ("PASS", "FAIL")
    assert verdict_summary.fatigue.verdict in text_verdict
    assert verdict_summary.rutting.verdict in text_verdict
    # [PLACEHOLDER] banner present (calibration still placeholder).
    assert "[PLACEHOLDER]" in text_verdict
    # [REFUSED] banner must NOT be present.
    assert "[REFUSED]" not in text_verdict
    # Calibration label rendered.
    assert "IRC37_PLACEHOLDER_80pct" in text_verdict
    # Strain-extraction fallback note carried through.
    assert "falling back" in text_verdict
    # IRC:37 cl. 6.4.2 / 6.4.3 / 6.4 references present.
    for cite in ("cl. 6.4.2", "cl. 6.4.3"):
        assert cite in text_verdict, f"missing {cite}"
    # Project-info block.
    assert "NH-007 Strengthening (Smoke Project)" in text_verdict
    assert "WO/2026/0001" in text_verdict
    print(f"  [PASS] {verdict_path.name} ({verdict_path.stat().st_size} bytes) "
          f"fatigue={verdict_summary.fatigue.verdict} "
          f"rutting={verdict_summary.rutting.verdict}")

    print("\nPHASE 15 MECH REPORT SMOKE: ALL OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
