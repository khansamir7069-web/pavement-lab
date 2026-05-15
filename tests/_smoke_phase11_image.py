"""Phase-11 image-wiring smoke.

Asserts the four step-3 deliverables in one headless run:

  1. Image pipeline round-trip — attach normalizes to JPEG q85 with
     max-edge 1600 px, content-addresses by SHA-256, dedupes on
     re-upload; list_evidence / delete_evidence / traversal-guard /
     delete_project_images all behave.
  2. Cascade — Database.delete_project removes both the survey row
     AND the project's image folder tree.
  3. Per-distress round-trip — attaching images to a specific
     distress row in ConditionSurveyPanel survives save + reload.
  4. Report embedding — build_condition_docx with survey-wide and
     per-distress images writes a docx whose 'Image Evidence'
     section contains the expected sub-headings AND embeds real
     JPEG image parts (verified via Document.part.related_parts).
"""
from __future__ import annotations

import importlib
import os
import sys
import tempfile
from pathlib import Path

# Redirect both the DB and IMAGES_DIR to a tempdir BEFORE any app.*
# module imports the config — modules like image_pipeline /
# condition_survey_panel / condition_report capture IMAGES_DIR at
# import time so we override and reload them.
os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
_tmp = Path(tempfile.mkdtemp())

import app.config as _cfg  # noqa: E402
_cfg.DB_PATH = _tmp / "phase11.db"
_cfg.IMAGES_DIR = _tmp / "images"
_cfg.IMAGES_DIR.mkdir(parents=True, exist_ok=True)

# Reload modules that captured IMAGES_DIR at import time.
from app.core.condition_survey import image_pipeline as _ip  # noqa: E402
importlib.reload(_ip)
from app.reports import condition_report as _cr  # noqa: E402
importlib.reload(_cr)
from app.ui.widgets import distress_images_dialog as _did  # noqa: E402
importlib.reload(_did)
from app.ui.widgets import condition_survey_panel as _csp  # noqa: E402
importlib.reload(_csp)

from docx import Document  # noqa: E402
from PIL import Image  # noqa: E402
from PySide6.QtWidgets import QApplication, QFileDialog, QMessageBox  # noqa: E402

QMessageBox.information = staticmethod(lambda *a, **k: 0)
QMessageBox.warning     = staticmethod(lambda *a, **k: 0)
QMessageBox.critical    = staticmethod(lambda *a, **k: 0)

from app.core import (  # noqa: E402
    ConditionSurveyInput,
    DistressRecord,
    compute_condition_survey,
)
from app.db.repository import Database  # noqa: E402
from app.reports import ConditionReportContext, build_condition_docx  # noqa: E402


def _synth_png(path: Path, size=(2400, 1800), color=(40, 160, 90)) -> Path:
    Image.new("RGB", size, color=color).save(path, format="PNG")
    return path


def _docx_text(p: Path) -> str:
    doc = Document(str(p))
    out = [pp.text for pp in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                out.append(c.text)
    return "\n".join(out)


def _docx_image_count(p: Path) -> int:
    doc = Document(str(p))
    return sum(
        1 for part in doc.part.related_parts.values()
        if getattr(part, "content_type", "").startswith("image/")
    )


def main() -> int:
    print("=== 1) Image pipeline round-trip ===")
    src1 = _synth_png(_tmp / "src1.png", (4000, 3000))
    ev = _ip.attach_image(project_id=11, survey_id=7, src_path=src1)
    assert ev.width_px == 1600 and ev.height_px == 1200, (ev.width_px, ev.height_px)
    assert ev.relative_path.startswith("condition/11/7/") and ev.relative_path.endswith(".jpg")
    assert "\\" not in ev.relative_path
    abs_path = _cfg.IMAGES_DIR / Path(ev.relative_path)
    assert abs_path.is_file()
    # Dedup
    ev2 = _ip.attach_image(11, 7, src1)
    assert ev2.relative_path == ev.relative_path
    # List / delete / traversal guard
    assert len(_ip.list_evidence(11, 7)) == 1
    assert _ip.delete_evidence(11, 7, "condition/99/7/whatever.jpg") is False
    assert _ip.delete_evidence(11, 7, ev.relative_path) is True
    assert _ip.list_evidence(11, 7) == ()
    # Project purge
    _ip.attach_image(11, 7, src1)
    _ip.attach_image(11, 8, src1)
    removed = _ip.delete_project_images(11)
    assert removed >= 2 and not (_cfg.IMAGES_DIR / "condition" / "11").exists()
    print(f"  [PASS] attach/dedup/list/delete/traversal/purge ({removed} purged)")

    print("\n=== 2) Database.delete_project cascade ===")
    db = Database(_cfg.DB_PATH)
    proj = db.create_project(work_name="Phase 11 cascade")
    _ip.attach_image(proj.id, 0, src1)
    _ip.attach_image(proj.id, 0, _synth_png(_tmp / "src2.png", (1200, 900),
                                            color=(200, 80, 40)))
    survey = ConditionSurveyInput(
        work_name=proj.work_name,
        records=(DistressRecord("cracking", "medium", length_m=80.0),),
    )
    db.save_condition_survey(project_id=proj.id,
                             result=compute_condition_survey(survey))
    assert db.latest_condition_survey(proj.id) is not None
    proj_dir = _cfg.IMAGES_DIR / "condition" / str(proj.id)
    assert proj_dir.is_dir()
    assert db.delete_project(proj.id) is True
    assert db.latest_condition_survey(proj.id) is None
    assert not proj_dir.exists(), "image folder not cleaned up on delete_project"
    print("  [PASS] DB row + image tree both gone")

    print("\n=== 3) Per-distress round-trip (panel save/reload) ===")
    app = QApplication.instance() or QApplication(sys.argv)
    from app.ui.main_window import MainWindow
    w = MainWindow()
    proj2 = w.db.create_project(work_name="Phase 11 per-distress")
    w._current_project_id = proj2.id
    w.condition.set_project(proj2.id, proj2.work_name)

    # Attach one image to row 0 by patching the dialog flow inline.
    w.condition._add_row(distress_code="potholes", severity="high", count=3)
    btn_img = w.condition.table.cellWidget(0, 6)
    assert btn_img is not None
    pothole_src = _synth_png(_tmp / "pothole.png", (1500, 1000),
                             color=(60, 60, 60))
    ev_row = _ip.attach_image(proj2.id, _csp.DRAFT_SURVEY_ID, pothole_src)
    btn_img._image_paths = [ev_row.relative_path]
    btn_img.setText(f"{len(btn_img._image_paths)} image(s)")

    # Survey-wide gallery image too (different file -> different sha).
    survey_src = _synth_png(_tmp / "survey_wide.png", (1800, 1200),
                            color=(160, 200, 60))
    ev_sw = _ip.attach_image(proj2.id, _csp.DRAFT_SURVEY_ID, survey_src)
    w.condition._evidence = [ev_sw.relative_path]
    w.condition._refresh_evidence_mirror()

    w.condition._on_compute()
    last = w.condition.last_result()
    assert last is not None
    assert last.inputs.records[0].image_paths == (ev_row.relative_path,)
    assert last.inputs.image_paths == (ev_sw.relative_path,)
    w.condition._on_save()

    # Reload from DB through set_project.
    w.condition.set_project(None)
    w.condition.set_project(proj2.id, proj2.work_name)
    btn_img_reloaded = w.condition.table.cellWidget(0, 6)
    assert btn_img_reloaded is not None
    assert list(btn_img_reloaded._image_paths) == [ev_row.relative_path], \
        btn_img_reloaded._image_paths
    assert w.condition._evidence == [ev_sw.relative_path]
    print("  [PASS] per-distress image_paths and survey-wide gallery round-trip")

    print("\n=== 4) Report embedding ===")
    out_doc = _tmp / "phase11_condition.docx"
    build_condition_docx(
        out_doc,
        ConditionReportContext(
            work_name=proj2.work_name,
            project_title=proj2.work_name,
        ),
        last,
    )
    txt = _docx_text(out_doc)
    assert "Image Evidence" in txt, "missing Image Evidence heading"
    assert "Survey-wide" in txt, "missing Survey-wide sub-heading"
    assert "Distress 1 - Potholes (High)" in txt, \
        "missing per-distress sub-heading"
    embedded = _docx_image_count(out_doc)
    # Expect at least 2 image parts (survey-wide + per-distress); python-docx
    # may also list theme images, so check >= 2 rather than ==.
    assert embedded >= 2, f"too few embedded image parts: {embedded}"
    print(f"  [PASS] Word doc embeds {embedded} image part(s) "
          f"({out_doc.stat().st_size} bytes)")

    print("\nPHASE 11 IMAGE SMOKE: ALL OK")
    return 0


if __name__ == "__main__":
    sys.exit(main())
