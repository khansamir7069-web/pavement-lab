"""Material Quantity — Word section + standalone docx builder (Phase 7).

Reuses ``_docx_common`` helpers and the Phase-6 ``CodeRef`` registry — no
new docx primitives, no new code-citation strings.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core import MaterialQuantityResult
from app.core.material_quantity import REFERENCES as ENGINE_REFS

from ._docx_common import (
    add_code_references,
    add_heading,
    add_kv_table,
    add_note,
    add_p,
    add_signature_block,
    add_table,
    new_portrait_document,
)


@dataclass(frozen=True, slots=True)
class MaterialQuantityReportContext:
    project_title: str = ""
    work_name: str = ""
    work_order_no: str = ""
    work_order_date: str = ""
    client: str = ""
    agency: str = ""
    submitted_by: str = ""
    lab_name: str = "Pavement Laboratory"
    report_date: str = field(default_factory=lambda: datetime.now().strftime("%d-%b-%Y"))


def write_material_quantity_section(
    doc: Document,
    ctx: MaterialQuantityReportContext,
    result: MaterialQuantityResult,
    *, include_header: bool = True,
) -> None:
    if include_header:
        add_heading(doc, "BILL OF MATERIAL QUANTITIES",
                    level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        if ctx.project_title:
            add_p(doc, ctx.project_title, bold=True, size=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(doc,
              "Quantities estimated per MoRTH Section 400 / 500 and IRC:111 "
              "default densities and spray rates. Job-Mix Formula and "
              "lab-verified densities supersede the defaults shown here.",
              italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(doc, f"{ctx.lab_name}  •  Report Date: {ctx.report_date}",
              size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_heading(doc, "Project Information", level=2)
        add_kv_table(doc, (
            ("Name of Work",     ctx.work_name),
            ("Work Order No.",   ctx.work_order_no),
            ("Work Order Date",  ctx.work_order_date),
            ("Client",           ctx.client),
            ("Agency",           ctx.agency),
            ("Submitted By",     ctx.submitted_by),
        ))

    # Per-layer breakdown
    add_heading(doc, "Layer-wise Breakdown", level=2)
    rows: list[list[str]] = []
    for lr in result.layers:
        inp = lr.inputs
        density = (inp.density_t_m3 if inp.density_t_m3 is not None
                   else "(default)")
        binder = (f"{inp.binder_pct:.2f}" if inp.binder_pct is not None
                  else ("—" if lr.category != "bituminous_mix" else "(default)"))
        spray = (f"{inp.spray_rate_kgm2:.3f}"
                 if inp.spray_rate_kgm2 is not None
                 else ("(default)" if lr.category == "sprayed_coat" else "—"))
        rows.append([
            inp.layer_type,
            f"{inp.length_m:g}", f"{inp.width_m:g}",
            f"{inp.thickness_mm:g}" if lr.category != "sprayed_coat" else "—",
            str(density) if isinstance(density, str) else f"{density:g}",
            binder,
            spray,
            f"{lr.area_m2:.1f}",
            f"{lr.layer_tonnage_t:.2f}",
            f"{lr.binder_tonnage_t:.2f}",
        ])
    add_table(doc,
        ["Layer", "L (m)", "W (m)", "t (mm)", "ρ (t/m³)",
         "Pb (%)", "Spray (kg/m²)", "Area (m²)", "Layer (t)", "Binder (t)"],
        rows,
    )

    # Totals
    add_heading(doc, "Totals", level=2)
    add_kv_table(doc, (
        ("Total Area",          f"{result.total_area_m2:.1f} m²"),
        ("Total Layer Mass",    f"{result.total_layer_tonnage_t:.2f} t"),
        ("Total Binder Mass",   f"{result.total_binder_tonnage_t:.2f} t"),
    ))

    # Citations — pulled from the engine module's REFERENCES tuple
    add_code_references(doc, ENGINE_REFS, heading="References")

    if result.notes:
        add_note(doc, result.notes)


def build_material_quantity_docx(
    out_path: Path,
    ctx: MaterialQuantityReportContext,
    result: MaterialQuantityResult,
) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = new_portrait_document()
    write_material_quantity_section(doc, ctx, result, include_header=True)
    add_signature_block(doc)
    doc.save(out_path)
    return out_path
