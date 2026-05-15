"""Maintenance / Rehabilitation — Word report.

Each sub-section cites its governing Indian Roads Congress code:

* BBD Overlay      → IRC:81-1997 "Guidelines for Strengthening of Flexible
                     Road Pavements using Benkelman Beam Deflection Technique"
                     (First Revision). Key clauses: 6.5 (T-correction),
                     6.6 (seasonal correction), 7.2 (characteristic
                     deflection), 8.2 (BM-equivalent overlay), Plate 1
                     (allowable deflection vs MSA), Table 6 (layer
                     equivalencies).
* Cold Mix         → IRC:SP:100-2014 "Use of Cold Mix Technology in
                     Construction and Maintenance of Roads using Bitumen
                     Emulsion". Marshall-based mix design (cl. 7);
                     Optimum Pre-Wetting Water Content (cl. 7.3).
* Micro-Surfacing  → IRC:SP:81 "Tentative Specifications for Slurry Seal
                     and Microsurfacing" (2008, revised 2014). Type I/II/III
                     gradations + WTAT / LWT / cohesion tests.

Public entry points:
  * ``write_overlay_section``, ``write_cold_mix_section``,
    ``write_micro_surfacing_section`` — append to an existing Document.
  * ``build_maintenance_docx`` — stand-alone docx containing any subset of
    the three sections.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core import ColdMixResult, MicroSurfacingResult, OverlayResult

from ._docx_common import (
    add_heading,
    add_kv_table,
    add_note,
    add_p,
    add_signature_block,
    add_table,
    new_portrait_document,
)


@dataclass(frozen=True, slots=True)
class MaintenanceReportContext:
    project_title: str = ""
    work_name: str = ""
    work_order_no: str = ""
    work_order_date: str = ""
    client: str = ""
    agency: str = ""
    submitted_by: str = ""
    lab_name: str = "Pavement Laboratory"
    report_date: str = field(default_factory=lambda: datetime.now().strftime("%d-%b-%Y"))


# ---------------------------------------------------------------------------
# Overlay (IRC:81-1997)
# ---------------------------------------------------------------------------

def write_overlay_section(doc: Document, ctx: MaintenanceReportContext,
                          result: OverlayResult, *, header_level: int = 2) -> None:
    add_heading(doc, "Overlay Design — BBD Method (IRC:81-1997)",
                level=header_level)
    add_p(doc, "Design basis: IRC:81-1997 — Guidelines for Strengthening of "
               "Flexible Road Pavements using the Benkelman Beam Deflection "
               "Technique (First Revision).",
          italic=True, size=10)

    inp = result.inputs

    # Inputs
    add_heading(doc, "Inputs", level=3)
    add_kv_table(doc, (
        ("Road Category",                                inp.road_category),
        ("Pavement Temperature at Testing",              f"{inp.pavement_temp_c:g} °C"),
        ("Existing Bituminous-Layer Thickness",          f"{inp.bituminous_thickness_mm:g} mm"),
        ("Seasonal Correction Factor (IRC:81 cl. 6.6)",  f"{inp.season_factor:g}"),
        ("Subgrade Type",                                 inp.subgrade_type),
        ("Cumulative Design Traffic",                     f"{inp.design_traffic_msa:g} MSA"),
        ("Number of Deflection Readings",                f"{result.n_readings}"),
    ))

    if result.n_readings == 0:
        add_p(doc, result.notes, size=10)
        return

    # Reading list
    add_heading(doc, "Measured Rebound Deflections", level=3)
    add_table(doc, ["#", "Deflection (mm)"],
              [[str(i + 1), f"{d:.3f}"] for i, d in enumerate(inp.deflections_mm)])

    # Calculation steps
    add_heading(doc, "Calculation Summary", level=3)
    add_p(doc, "IRC:81-1997 cl. 6.5 — Temperature correction (to 35 °C):",
          size=10)
    add_p(doc, "   D_corrected = D_measured + (35 − T) × 0.0065 mm",
          italic=True, size=10)
    add_p(doc, "IRC:81-1997 cl. 6.6 — Seasonal correction factor applied "
               "uniformly to the corrected readings.", size=10)
    add_p(doc, "IRC:81-1997 cl. 7.2 — Characteristic deflection:", size=10)
    add_p(doc, "   D_c = mean + k · SD   "
               "(k = 2.0 for NH/SH/Expressway; k = 1.0 for other roads)",
          italic=True, size=10)
    add_p(doc, "IRC:81-1997 Plate 1 — Allowable rebound deflection vs MSA "
               "(log-linear interpolation between tabulated points).", size=10)
    add_p(doc, "IRC:81-1997 cl. 8.2 — BM-equivalent overlay thickness:", size=10)
    add_p(doc, "   h_BM = 550 × log10(D_c / D_allow) mm",
          italic=True, size=10)

    add_table(doc, ["Parameter", "Value", "Unit"], [
        ["Mean deflection",                f"{result.mean_deflection_mm:.4f}", "mm"],
        ["Standard deviation",             f"{result.stdev_deflection_mm:.4f}", "mm"],
        ["Characteristic deflection D_c",  f"{result.characteristic_deflection_mm:.4f}", "mm"],
        ["Allowable deflection D_allow",   f"{result.allowable_deflection_mm:.4f}", "mm"],
        ["BM-equivalent overlay thickness",
         f"{result.overlay_thickness_mm:.1f}", "mm"],
    ])

    if result.overlay_required:
        add_p(doc, f"Result: Overlay required. Provide {result.overlay_thickness_mm:.0f} mm "
                   "of BM-equivalent thickness, then convert to DBM / BC / GSB "
                   "courses using IRC:81-1997 Table 6 layer-equivalencies.",
              bold=True, size=11)
    else:
        add_p(doc, "Result: No overlay required (D_c ≤ D_allow).",
              bold=True, size=11)

    add_note(doc, result.notes or "")


# ---------------------------------------------------------------------------
# Cold Mix (IRC:SP:100-2014)
# ---------------------------------------------------------------------------

def write_cold_mix_section(doc: Document, ctx: MaintenanceReportContext,
                           result: ColdMixResult, *, header_level: int = 2) -> None:
    add_heading(doc, "Cold Bituminous Mix Design (IRC:SP:100-2014)",
                level=header_level)
    add_p(doc, "Design basis: IRC:SP:100-2014 — Use of Cold Mix Technology "
               "in Construction and Maintenance of Roads Using Bitumen "
               "Emulsion. Marshall-based mix design (cl. 7).",
          italic=True, size=10)

    inp = result.inputs
    add_heading(doc, "Inputs", level=3)
    add_kv_table(doc, (
        ("Mix Type / Gradation",                  inp.mix_type),
        ("Aggregate Mass (basis)",                f"{inp.aggregate_mass_kg:g} kg"),
        ("Bitumen-Emulsion Content (% of agg.)",  f"{inp.emulsion_pct:g} %"),
        ("Emulsion Residue (per datasheet)",      f"{inp.emulsion_residue_pct:g} %"),
        ("Pre-Wetting / Added Water (% of agg.)", f"{inp.water_addition_pct:g} %"),
        ("Mineral Filler (% of agg.)",            f"{inp.filler_pct:g} %"),
    ))

    add_heading(doc, "Computed Proportions (per basis aggregate mass)", level=3)
    add_table(doc, ["Component", "Mass (kg)", "% of aggregate"],
              [[c.name, f"{c.mass_kg:.2f}", f"{c.pct_of_aggregate:.2f}"]
               for c in result.components])

    lo, hi = result.spec_window_pct
    add_heading(doc, "Specification Check", level=3)
    add_table(doc, ["Property", "Value", "IRC:SP:100 Window", "Status"], [
        ["Residual Binder (% of aggregate)",
         f"{result.residual_binder_pct:.2f}",
         f"{lo:.1f} – {hi:.1f} %",
         "PASS" if result.pass_check else "FAIL"],
        ["Total Mix Mass", f"{result.total_mix_mass_kg:.2f} kg", "—", "—"],
    ])
    if not result.pass_check and result.pass_reasons:
        for reason in result.pass_reasons:
            add_p(doc, f"• {reason}", size=10)

    add_note(doc, result.notes or "")


# ---------------------------------------------------------------------------
# Micro-surfacing (IRC:SP:81)
# ---------------------------------------------------------------------------

def write_micro_surfacing_section(doc: Document, ctx: MaintenanceReportContext,
                                  result: MicroSurfacingResult,
                                  *, header_level: int = 2) -> None:
    add_heading(doc, "Micro-Surfacing Mix Design (IRC:SP:81)",
                level=header_level)
    add_p(doc, "Design basis: IRC:SP:81 — Tentative Specifications for "
               "Slurry Seal and Microsurfacing (2008, revised 2014).",
          italic=True, size=10)

    inp = result.inputs
    add_heading(doc, "Inputs", level=3)
    add_kv_table(doc, (
        ("Surfacing Type",                                 inp.surfacing_type),
        ("Aggregate Mass (basis)",                         f"{inp.aggregate_mass_kg:g} kg"),
        ("Polymer-Modified Emulsion (% of agg.)",          f"{inp.emulsion_pct:g} %"),
        ("Emulsion Residue (per datasheet)",               f"{inp.emulsion_residue_pct:g} %"),
        ("Additive Water (% of agg.)",                     f"{inp.additive_water_pct:g} %"),
        ("Mineral Filler — OPC (% of agg.)",               f"{inp.mineral_filler_pct:g} %"),
    ))

    add_heading(doc, "Computed Proportions (per basis aggregate mass)", level=3)
    add_table(doc, ["Component", "Mass (kg)", "% of aggregate"],
              [[c.name, f"{c.mass_kg:.2f}", f"{c.pct_of_aggregate:.2f}"]
               for c in result.components])

    rb_lo, rb_hi = result.spec_residual_binder_pct
    fl_lo, fl_hi = result.spec_filler_pct
    add_heading(doc, "Specification Check", level=3)
    add_table(doc, ["Property", "Value", f"IRC:SP:81 {inp.surfacing_type}",
                    "Status"], [
        ["Residual Binder (% of agg.)",
         f"{result.residual_binder_pct:.2f}",
         f"{rb_lo:.1f} – {rb_hi:.1f} %",
         "PASS" if not any('Residual binder' in r for r in result.pass_reasons) else "FAIL"],
        ["Additive Water (% of agg.)",
         f"{inp.additive_water_pct:.2f}",
         "5.0 – 12.0 %",
         "PASS" if not any('Additive water' in r for r in result.pass_reasons) else "FAIL"],
        ["Mineral Filler (% of agg.)",
         f"{inp.mineral_filler_pct:.2f}",
         f"{fl_lo:.1f} – {fl_hi:.1f} %",
         "PASS" if not any('Mineral filler' in r for r in result.pass_reasons) else "FAIL"],
        ["Total Water Demand (% of agg.)",
         f"{result.total_water_demand_pct:.2f}", "—", "—"],
    ])
    if not result.pass_check and result.pass_reasons:
        for reason in result.pass_reasons:
            add_p(doc, f"• {reason}", size=10)

    add_note(doc, result.notes or "")


# ---------------------------------------------------------------------------
# Stand-alone maintenance document
# ---------------------------------------------------------------------------

def build_maintenance_docx(
    out_path: Path,
    ctx: MaintenanceReportContext,
    *,
    overlay: Optional[OverlayResult] = None,
    cold_mix: Optional[ColdMixResult] = None,
    micro_surfacing: Optional[MicroSurfacingResult] = None,
) -> Path:
    """Build a stand-alone Word document with whichever sub-modules are given."""
    if not any((overlay, cold_mix, micro_surfacing)):
        raise ValueError("At least one maintenance result must be provided.")

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = new_portrait_document()

    add_heading(doc, "REPORT ON MAINTENANCE / REHABILITATION DESIGN",
                level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    if ctx.project_title:
        add_p(doc, ctx.project_title, bold=True, size=12,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc,
          "Design references: IRC:81-1997 (BBD), IRC:SP:100-2014 (Cold Mix), "
          "IRC:SP:81 (Micro-Surfacing).",
          italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"{ctx.lab_name}  •  Report Date: {ctx.report_date}",
          size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "Project Information", level=2)
    add_kv_table(doc, (
        ("Name of Work",     ctx.work_name),
        ("Work Order No.",   ctx.work_order_no),
        ("Work Order Date",  ctx.work_order_date),
        ("Client",           ctx.client),
        ("Agency",           ctx.agency),
        ("Submitted By",     ctx.submitted_by),
    ))

    if overlay is not None:
        write_overlay_section(doc, ctx, overlay)
    if cold_mix is not None:
        write_cold_mix_section(doc, ctx, cold_mix)
    if micro_surfacing is not None:
        write_micro_surfacing_section(doc, ctx, micro_surfacing)

    add_signature_block(doc)
    doc.save(out_path)
    return out_path
