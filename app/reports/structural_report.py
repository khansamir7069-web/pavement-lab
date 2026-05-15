"""Flexible Pavement Structural Design — Word report.

Authoritative basis (cited in the report itself):
    IRC:37-2018 — "Guidelines for the Design of Flexible Pavements"
    (Indian Roads Congress, 4th Revision).
    * cl. 4.6 — Cumulative design traffic in MSA:
          N = (365·A·((1+r)^n − 1)/r) · D · F · 10⁻⁶
    * Annex E — Subgrade resilient modulus from CBR:
          M_R (MPa) = 10·CBR             for CBR ≤ 5 %
                    = 17.6·CBR^0.64      for CBR > 5 %
    * IRC:37-2018 Plates 1–4 / catalogue — typical layer compositions for
      flexible pavements. The composition produced by this software is a
      catalogue-style suggestion and must be cross-checked against the
      relevant Plate or verified with IITPAVE before adoption.

This file produces ONE Word section (it can be appended to the combined
report builder via ``write_structural_section``) OR a stand-alone docx
via ``build_structural_docx``.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Mapping, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core import StructuralResult

from ._docx_common import (
    add_heading,
    add_kv_table,
    add_note,
    add_p,
    add_signature_block,
    add_table,
    new_portrait_document,
)


@dataclass(frozen=True, slots=True)
class StructuralReportContext:
    project_title: str = ""
    work_name: str = ""
    work_order_no: str = ""
    work_order_date: str = ""
    client: str = ""
    agency: str = ""
    submitted_by: str = ""
    lab_name: str = "Pavement Laboratory"
    report_date: str = field(default_factory=lambda: datetime.now().strftime("%d-%b-%Y"))


# ---------------------------------------------------------------------------
# Section writer (re-used by combined report)
# ---------------------------------------------------------------------------

def write_structural_section(
    doc: Document, ctx: StructuralReportContext, result: StructuralResult,
    *, include_header: bool = True,
) -> None:
    """Append the IRC:37 structural design section to ``doc`` in place."""
    if include_header:
        add_heading(doc, "FLEXIBLE PAVEMENT STRUCTURAL DESIGN",
                    level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        if ctx.project_title:
            add_p(doc, ctx.project_title, bold=True, size=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(doc, "Design basis: IRC:37-2018 — Guidelines for the Design of "
                   "Flexible Pavements (4th Revision).",
              size=10, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
        add_p(doc, f"{ctx.lab_name}  •  Report Date: {ctx.report_date}",
              size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Project information block
    add_heading(doc, "1. Project Information", level=2)
    add_kv_table(doc, (
        ("Name of Work",     ctx.work_name),
        ("Work Order No.",   ctx.work_order_no),
        ("Work Order Date",  ctx.work_order_date),
        ("Client",           ctx.client),
        ("Agency",           ctx.agency),
        ("Submitted By",     ctx.submitted_by),
    ))

    # Inputs
    add_heading(doc, "2. Design Inputs", level=2)
    inp = result.inputs
    add_kv_table(doc, (
        ("Road Category",                        inp.road_category),
        ("Design Life (years)",                  f"{inp.design_life_years}"),
        ("Initial Commercial Traffic (CVPD)",    f"{inp.initial_cvpd:g}"),
        ("Traffic Growth Rate (% / year)",       f"{inp.growth_rate_pct:g}"),
        ("Vehicle Damage Factor (VDF)",          f"{inp.vdf:g}"),
        ("Lane Distribution Factor (LDF)",       f"{inp.ldf:g}"),
        ("Subgrade CBR — 4-day soaked (%)",      f"{inp.subgrade_cbr_pct:g}"),
        ("Resilient Modulus (user-supplied)",
            (f"{inp.resilient_modulus_mpa:g} MPa"
             if inp.resilient_modulus_mpa else "auto (from CBR)")),
    ))

    # Computed traffic
    add_heading(doc, "3. Cumulative Design Traffic", level=2)
    add_p(doc, "Per IRC:37-2018 cl. 4.6:")
    add_p(doc, "N = (365 · A · [(1 + r)^n − 1] / r) · D · F · 10⁻⁶",
          italic=True, size=11)
    add_table(doc, ["Parameter", "Symbol", "Value", "Unit"], [
        ["Initial Commercial Traffic", "A", f"{inp.initial_cvpd:g}", "CVPD"],
        ["Annual Growth Rate",         "r", f"{inp.growth_rate_pct/100:.4f}", "—"],
        ["Design Life",                "n", f"{inp.design_life_years}", "years"],
        ["Lane Distribution Factor",   "D", f"{inp.ldf:g}", "—"],
        ["Vehicle Damage Factor",      "F", f"{inp.vdf:g}", "—"],
        ["Growth Factor [(1+r)^n−1]/r", "—", f"{result.growth_factor:.4f}", "—"],
        ["Cumulative Design Traffic",  "N", f"{result.design_msa:.2f}", "MSA"],
    ])

    # Subgrade modulus
    add_heading(doc, "4. Subgrade Resilient Modulus", level=2)
    add_p(doc, "Per IRC:37-2018 Annex E:")
    add_p(doc, "  M_R = 10 · CBR              for CBR ≤ 5 %", italic=True, size=11)
    add_p(doc, "  M_R = 17.6 · CBR^0.64       for CBR > 5 %", italic=True, size=11)
    add_table(doc, ["Property", "Value", "Unit"], [
        ["CBR (4-day soaked)",                  f"{inp.subgrade_cbr_pct:g}", "%"],
        ["Resilient Modulus (M_R)",             f"{result.subgrade_mr_mpa:.1f}", "MPa"],
    ])

    # Layer composition
    add_heading(doc, "5. Suggested Pavement Composition", level=2)
    rows = [
        [ly.name, ly.material or "—",
         f"{ly.thickness_mm:.0f}",
         f"{ly.modulus_mpa:.0f}" if ly.modulus_mpa else "—"]
        for ly in result.composition
    ]
    rows.append(["TOTAL", "—", f"{result.total_pavement_thickness_mm:.0f}", "—"])
    add_table(doc, ["Layer", "Material", "Thickness (mm)",
                    "Typical Modulus (MPa)"], rows)
    add_note(doc,
        "Layer composition is a catalogue-style suggestion. Cross-check "
        "against IRC:37-2018 Plates 1–4 and run a mechanistic analysis "
        "(IITPAVE) for the fatigue and rutting limits in cl. 6.4 before "
        "adoption for construction."
    )

    # Mechanistic checks (placeholder until IITPAVE is integrated)
    add_heading(doc, "6. Mechanistic Checks (IITPAVE)", level=2)
    add_table(doc, ["Check", "Status"], [
        ["Bituminous-layer fatigue (IRC:37-2018 cl. 6.4.1)",
         result.fatigue_check or "Not performed."],
        ["Subgrade rutting (IRC:37-2018 cl. 6.4.2)",
         result.rutting_check or "Not performed."],
    ])

    # Notes
    if result.notes:
        add_heading(doc, "7. Designer Notes", level=2)
        add_p(doc, result.notes, size=10)

    add_note(doc,
        "References: IRC:37-2018 'Guidelines for the Design of Flexible "
        "Pavements' (Indian Roads Congress, 4th Revision)."
    )


# ---------------------------------------------------------------------------
# Stand-alone document
# ---------------------------------------------------------------------------

def build_structural_docx(
    out_path: Path, ctx: StructuralReportContext, result: StructuralResult,
) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = new_portrait_document()
    write_structural_section(doc, ctx, result, include_header=True)
    add_signature_block(doc)
    doc.save(out_path)
    return out_path
