"""Mechanistic Validation — Word section + standalone docx (Phase 15 — Priority 1).

Consumes :class:`app.core.MechanisticValidationSummary` (Phase 14) and
renders an IRC:37-2018 cl. 6.4 fatigue + rutting verdict report.

Refusal-aware: when the summary's refusal gate fired (placeholder
mechanistic input, missing strain or missing BC modulus) the report
emits a prominent REFUSED banner instead of a verdict KV table — the
underlying numbers are intentionally absent in that case (Phase 14
contract). When verdicts are present but calibration constants are
still the IRC37_PLACEHOLDER defaults, a softer PLACEHOLDER banner is
rendered so the reader cannot mistake the values for field-calibrated
output.

Additive only. Reuses Phase-6 ``_docx_common`` helpers exclusively.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core import (
    FatigueCheck,
    MechanisticValidationSummary,
    RuttingCheck,
)
from app.core.mechanistic_validation import REFERENCES as ENGINE_REFS

from ._docx_common import (
    add_code_references,
    add_heading,
    add_kv_table,
    add_note,
    add_p,
    add_placeholder_banner,
    add_signature_block,
    add_table,
    new_portrait_document,
)


# ---------------------------------------------------------------------------
# Context
# ---------------------------------------------------------------------------

@dataclass(frozen=True, slots=True)
class MechanisticReportContext:
    project_title: str = ""
    work_name: str = ""
    work_order_no: str = ""
    work_order_date: str = ""
    client: str = ""
    agency: str = ""
    submitted_by: str = ""
    lab_name: str = "Pavement Laboratory"
    report_date: str = field(default_factory=lambda: datetime.now().strftime("%d-%b-%Y"))


# ---------------------------------------------------------------------------
# Formatters
# ---------------------------------------------------------------------------

_DASH = "—"


def _fmt_optional_float(value: float | None, fmt: str = "{:.2f}") -> str:
    return _DASH if value is None else fmt.format(value)


def _fmt_verdict(verdict: str | None) -> str:
    return _DASH if verdict is None else verdict


# ---------------------------------------------------------------------------
# Per-check renderers
# ---------------------------------------------------------------------------

def _render_fatigue(doc: Document, check: FatigueCheck) -> None:
    add_heading(doc, "Fatigue (Bottom of Bituminous-Bound Layer)", level=2)
    add_p(doc,
          "Model (IRC:37-2018 cl. 6.4.2):  "
          "N_f = C · k1 · (1/ε_t)^k2 · (1/E)^k3   →   N_f_MSA = N_f / 10^6",
          italic=True, size=10)
    add_kv_table(doc, (
        ("Tensile strain ε_t (microstrain)",
            _fmt_optional_float(check.epsilon_t_microstrain, "{:.2f}")),
        ("Bituminous modulus E_BC (MPa)",
            _fmt_optional_float(check.e_bc_mpa, "{:.0f}")),
        ("Design traffic (MSA)",
            f"{check.design_msa:.2f}"),
        ("C factor (VBE/Va adjustment)",
            f"{check.c_factor:g}"),
        ("Cumulative fatigue life (MSA)",
            _fmt_optional_float(check.cumulative_life_msa, "{:.2f}")),
        ("Verdict",
            _fmt_verdict(check.verdict)),
    ))
    add_heading(doc, "Fatigue Calibration", level=3)
    cal = check.calibration
    add_table(doc, ["Constant", "Symbol", "Value"], [
        ["Calibration label",     "—",  cal.label],
        ["Reliability",           "—",  f"{cal.reliability_pct}%"],
        ["Coefficient",           "k1", f"{cal.k1:.4e}"],
        ["Strain exponent",       "k2", f"{cal.k2:g}"],
        ["Modulus exponent",      "k3", f"{cal.k3:g}"],
        ["Placeholder",           "—",  "Yes" if cal.is_placeholder else "No"],
    ])
    if check.refused:
        add_placeholder_banner(doc,
            f"[REFUSED — fatigue] {check.refused_reason}"
        )
    if check.notes:
        add_note(doc, check.notes)


def _render_rutting(doc: Document, check: RuttingCheck) -> None:
    add_heading(doc, "Rutting (Top of Subgrade)", level=2)
    add_p(doc,
          "Model (IRC:37-2018 cl. 6.4.3):  "
          "N_r = k_r · (1/ε_v)^k_v   →   N_r_MSA = N_r / 10^6",
          italic=True, size=10)
    add_kv_table(doc, (
        ("Vertical strain ε_v (microstrain)",
            _fmt_optional_float(check.epsilon_v_microstrain, "{:.2f}")),
        ("Design traffic (MSA)",
            f"{check.design_msa:.2f}"),
        ("Cumulative rutting life (MSA)",
            _fmt_optional_float(check.cumulative_life_msa, "{:.2f}")),
        ("Verdict",
            _fmt_verdict(check.verdict)),
    ))
    add_heading(doc, "Rutting Calibration", level=3)
    cal = check.calibration
    add_table(doc, ["Constant", "Symbol", "Value"], [
        ["Calibration label",     "—",   cal.label],
        ["Reliability",           "—",   f"{cal.reliability_pct}%"],
        ["Coefficient",           "k_r", f"{cal.k_r:.4e}"],
        ["Strain exponent",       "k_v", f"{cal.k_v:g}"],
        ["Placeholder",           "—",   "Yes" if cal.is_placeholder else "No"],
    ])
    if check.refused:
        add_placeholder_banner(doc,
            f"[REFUSED — rutting] {check.refused_reason}"
        )
    if check.notes:
        add_note(doc, check.notes)


# ---------------------------------------------------------------------------
# Section + standalone builder
# ---------------------------------------------------------------------------

def write_mechanistic_section(
    doc: Document,
    ctx: MechanisticReportContext,
    summary: MechanisticValidationSummary,
    *,
    include_header: bool = True,
) -> None:
    if include_header:
        add_heading(doc, "MECHANISTIC VALIDATION (IRC:37-2018 cl. 6.4)",
                    level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        if ctx.project_title:
            add_p(doc, ctx.project_title, bold=True, size=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(doc,
              "Elastic-layer analysis (IITPAVE / IRC:37-2018 cl. 6.2) "
              "followed by fatigue (cl. 6.4.2) and rutting (cl. 6.4.3) "
              "life checks. Verdicts are refused when the upstream "
              "mechanistic result is flagged placeholder.",
              italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(doc, f"{ctx.lab_name}  -  Report Date: {ctx.report_date}",
              size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_heading(doc, "Project Information", level=2)
        add_kv_table(doc, (
            ("Name of Work",    ctx.work_name),
            ("Work Order No.",  ctx.work_order_no),
            ("Client",          ctx.client),
            ("Agency",          ctx.agency),
            ("Submitted By",    ctx.submitted_by),
        ))

    # --- Overall refusal banner (Phase 14 safety contract) --------------
    if summary.refused:
        add_placeholder_banner(doc,
            f"[REFUSED] {summary.refused_reason or 'Final verdict refused.'}"
        )
    elif summary.is_placeholder:
        # Verdicts produced but calibration constants are placeholder.
        add_placeholder_banner(doc, "[PLACEHOLDER] " + summary.notes)

    # --- Overall summary KV --------------------------------------------
    add_heading(doc, "Validation Summary", level=2)
    add_kv_table(doc, (
        ("Refused",            "Yes" if summary.refused else "No"),
        ("Placeholder",        "Yes" if summary.is_placeholder else "No"),
        ("Fatigue verdict",    _fmt_verdict(summary.fatigue.verdict)),
        ("Rutting verdict",    _fmt_verdict(summary.rutting.verdict)),
        ("Fatigue life (MSA)",
            _fmt_optional_float(summary.fatigue.cumulative_life_msa, "{:.2f}")),
        ("Rutting life (MSA)",
            _fmt_optional_float(summary.rutting.cumulative_life_msa, "{:.2f}")),
        ("Design traffic (MSA)",
            f"{summary.fatigue.design_msa:.2f}"),
    ))

    # --- Per-check sections --------------------------------------------
    _render_fatigue(doc, summary.fatigue)
    _render_rutting(doc, summary.rutting)

    # --- Notes (strain-extraction fallback warnings etc.) --------------
    if summary.notes:
        add_heading(doc, "Notes", level=3)
        add_note(doc, summary.notes)

    # --- References (IRC:37 cl. 6.2 / 6.4.2 / 6.4.3 / 6.4) -------------
    add_code_references(doc, ENGINE_REFS, heading="References")


def build_mechanistic_docx(
    out_path: Path,
    ctx: MechanisticReportContext,
    summary: MechanisticValidationSummary,
) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = new_portrait_document()
    write_mechanistic_section(doc, ctx, summary, include_header=True)
    add_signature_block(doc)
    doc.save(out_path)
    return out_path
