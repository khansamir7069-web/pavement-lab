"""Rehabilitation Recommendations — Word section + standalone docx
(Phase 15 — Priority 2).

Consumes a Phase-12 :class:`app.core.RehabSynthesisResult` and renders
the survey-level treatment recommendations as a prioritized table plus
a per-recommendation explainability block (reason / triggers /
next-module / source-tagged references). The Word output is
intentionally explainable: every recommendation tells the reader *why*
it fired, not just *what* it is.

Refusal vs placeholder behaviour:
  * ``synthesis.is_placeholder=True``    -> [PLACEHOLDER] banner with
                                            the engine's calibration
                                            caveat.
  * No recommendations triggered         -> a single italic note ("No
                                            treatment recommendations
                                            triggered for this
                                            survey.") — this is a
                                            valid outcome for an
                                            Excellent-PCI / no-distress
                                            survey-routine.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core import (
    REHAB_REFERENCES,
    RehabSynthesisResult,
    TreatmentRecommendation,
)

from ._docx_common import (
    add_code_references,
    add_heading,
    add_kv_table,
    add_note,
    add_p,
    add_placeholder_banner,
    add_signature_block,
    add_table,
    new_portrait_document,
)


# ---------------------------------------------------------------------------
# Context
# ---------------------------------------------------------------------------

@dataclass(frozen=True, slots=True)
class RehabReportContext:
    project_title: str = ""
    work_name: str = ""
    work_order_no: str = ""
    work_order_date: str = ""
    client: str = ""
    agency: str = ""
    submitted_by: str = ""
    lab_name: str = "Pavement Laboratory"
    report_date: str = field(default_factory=lambda: datetime.now().strftime("%d-%b-%Y"))


# ---------------------------------------------------------------------------
# Formatters
# ---------------------------------------------------------------------------

def _refs_inline(rec: TreatmentRecommendation) -> str:
    """Join the recommendation's reference code_ids inline (table cell)."""
    return ", ".join(ref.code_id for ref in rec.references) or "—"


def _triggers_inline(rec: TreatmentRecommendation) -> str:
    return ", ".join(rec.triggers) or "—"


def _next_module_label(rec: TreatmentRecommendation) -> str:
    return rec.next_module if rec.next_module and rec.next_module != "none" else "—"


# ---------------------------------------------------------------------------
# Section + standalone builder
# ---------------------------------------------------------------------------

def write_rehab_section(
    doc: Document,
    ctx: RehabReportContext,
    synthesis: RehabSynthesisResult,
    *,
    include_header: bool = True,
) -> None:
    if include_header:
        add_heading(doc, "REHABILITATION RECOMMENDATIONS",
                    level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        if ctx.project_title:
            add_p(doc, ctx.project_title, bold=True, size=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(doc,
              "Survey-level treatment synthesis derived from the "
              "pavement condition survey (PCI + distress breakdown) "
              "and — where available — the project's traffic / "
              "maintenance design data. Recommendation thresholds are "
              "IRC:82 / IRC:81 / IRC:115 / IRC:SP:81 / IRC:SP:101 "
              "shape but remain uncalibrated placeholder values.",
              italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(doc, f"{ctx.lab_name}  -  Report Date: {ctx.report_date}",
              size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_heading(doc, "Project Information", level=2)
        add_kv_table(doc, (
            ("Name of Work",    ctx.work_name),
            ("Work Order No.",  ctx.work_order_no),
            ("Client",          ctx.client),
            ("Agency",          ctx.agency),
            ("Submitted By",    ctx.submitted_by),
        ))

    # --- Placeholder banner --------------------------------------------
    if synthesis.is_placeholder:
        add_placeholder_banner(doc, "[PLACEHOLDER] " + (synthesis.notes or ""))

    # --- Context summary -----------------------------------------------
    add_heading(doc, "Synthesis Summary", level=2)
    add_kv_table(doc, (
        ("Context",            synthesis.context_summary),
        ("Recommendations",    str(len(synthesis.recommendations))),
        ("Placeholder",        "Yes" if synthesis.is_placeholder else "No"),
    ))

    # --- Prioritized recommendations table -----------------------------
    add_heading(doc, "Prioritized Treatment Recommendations", level=2)
    if not synthesis.recommendations:
        add_p(doc,
              "No treatment recommendations triggered for this survey.",
              italic=True, size=10)
    else:
        rows: list[list[str]] = []
        for rec in synthesis.recommendations:
            rows.append([
                str(rec.priority),
                rec.label,
                _triggers_inline(rec),
                _next_module_label(rec),
                _refs_inline(rec),
            ])
        add_table(doc, [
            "Priority", "Treatment", "Triggers",
            "Next Module", "References",
        ], rows)

    # --- Per-recommendation explainability blocks ----------------------
    if synthesis.recommendations:
        add_heading(doc, "Recommendation Detail", level=2)
        for idx, rec in enumerate(synthesis.recommendations, start=1):
            add_heading(doc,
                        f"#{idx}  {rec.label}  (priority {rec.priority})",
                        level=3)
            add_kv_table(doc, (
                ("Category",       rec.category),
                ("Reason",         rec.reason),
                ("Triggers",       _triggers_inline(rec)),
                ("Next module",    _next_module_label(rec)),
                ("References",     _refs_inline(rec)),
                ("Placeholder",    "Yes" if rec.is_placeholder else "No"),
            ))

    # --- Engine-level notes --------------------------------------------
    if synthesis.notes:
        add_heading(doc, "Notes", level=3)
        add_note(doc, synthesis.notes)

    # --- References block ----------------------------------------------
    # Prefer the synthesis-bundled references (IRC:82 / IRC:81 / etc.) so
    # an empty-recommendation report still cites the governing codes.
    refs = synthesis.references or REHAB_REFERENCES
    add_code_references(doc, refs, heading="References")


def build_rehab_docx(
    out_path: Path,
    ctx: RehabReportContext,
    synthesis: RehabSynthesisResult,
) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = new_portrait_document()
    write_rehab_section(doc, ctx, synthesis, include_header=True)
    add_signature_block(doc)
    doc.save(out_path)
    return out_path
