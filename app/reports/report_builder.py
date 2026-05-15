"""Module-aware combined report builder (Phase 6).

Discovers which modules of a project have saved data and produces a single
Word document containing every available section, each with its governing
IRC code citations:

    * Bituminous Mix Design     → MoRTH Section 500 / IRC:111 (via the
                                  existing ``word_report.build_mix_design_docx``
                                  helpers, only when a live ``MixDesignResult``
                                  is supplied — the DB stores a summary, not
                                  the full re-compute payload).
    * Flexible Pavement Design  → IRC:37-2018
    * BBD Overlay               → IRC:81-1997
    * Cold Mix                  → IRC:SP:100-2014
    * Micro-Surfacing           → IRC:SP:81

Inputs that are persisted in the DB (structural, maintenance) are
re-hydrated via the deterministic engine entry points
(``compute_structural_design``, ``compute_overlay``, ``compute_cold_mix``,
``compute_micro_surfacing``) so the report reproduces exactly what the
user saw at save time. No state escapes the engine.
"""
from __future__ import annotations

import json
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Mapping, Optional

from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core import (
    ColdMixInput,
    ConditionSurveyInput,
    DistressRecord,
    LayerInput,
    MaterialQuantityInput,
    MicroSurfacingInput,
    MixDesignResult,
    OverlayInput,
    StructuralInput,
    TrafficInput,
    compute_cold_mix,
    compute_condition_survey,
    compute_material_quantity,
    compute_micro_surfacing,
    compute_overlay,
    compute_structural_design,
    compute_traffic_analysis,
)
from app.core.import_summary import ImportedMixResult
from app.graphs import MarshallChartSet, build_chart_set

from ._docx_common import (
    add_heading,
    add_kv_table,
    add_p,
    add_signature_block,
    new_portrait_document,
)
from .maintenance_report import (
    MaintenanceReportContext,
    write_cold_mix_section,
    write_micro_surfacing_section,
    write_overlay_section,
)
from .material_qty_report import (
    MaterialQuantityReportContext,
    write_material_quantity_section,
)
from .traffic_report import (
    TrafficReportContext,
    write_traffic_section,
)
from .condition_report import (
    ConditionReportContext,
    write_condition_section,
)
from .rehab_report import (
    RehabReportContext,
    write_rehab_section,
)
from .structural_report import (
    StructuralReportContext,
    write_structural_section,
)


# ---------------------------------------------------------------------------
# Combined-report context
# ---------------------------------------------------------------------------

@dataclass(frozen=True, slots=True)
class CombinedReportContext:
    project_title: str = ""
    work_name: str = ""
    work_order_no: str = ""
    work_order_date: str = ""
    client: str = ""
    agency: str = ""
    submitted_by: str = ""
    lab_name: str = "Pavement Laboratory"
    report_date: str = field(default_factory=lambda: datetime.now().strftime("%d-%b-%Y"))
    binder_grade: str = ""
    binder_properties: Mapping[str, float] = field(default_factory=dict)
    mix_type_key: str = ""


# ---------------------------------------------------------------------------
# Re-hydrate persisted module rows via the deterministic engines
# ---------------------------------------------------------------------------

def _coerce_tuple(v) -> tuple:
    return tuple(v) if isinstance(v, list) else v


def _rehydrate_structural(sd_row) -> "StructuralResult | None":
    if not sd_row or not sd_row.inputs_json:
        return None
    try:
        d = json.loads(sd_row.inputs_json)
    except json.JSONDecodeError:
        return None
    inp = StructuralInput(
        road_category=d.get("road_category", "NH / SH"),
        design_life_years=int(d.get("design_life_years", 15)),
        initial_cvpd=float(d.get("initial_cvpd", 2000.0)),
        growth_rate_pct=float(d.get("growth_rate_pct", 7.5)),
        vdf=float(d.get("vdf", 2.5)),
        ldf=float(d.get("ldf", 0.75)),
        subgrade_cbr_pct=float(d.get("subgrade_cbr_pct", 5.0)),
        resilient_modulus_mpa=d.get("resilient_modulus_mpa"),
        notes=d.get("notes", "") or "",
    )
    return compute_structural_design(inp)


def _rehydrate_overlay(row) -> "OverlayResult | None":
    if not row or not row.inputs_json:
        return None
    try:
        d = json.loads(row.inputs_json)
    except json.JSONDecodeError:
        return None
    inp = OverlayInput(
        deflections_mm=tuple(d.get("deflections_mm") or ()),
        pavement_temp_c=float(d.get("pavement_temp_c", 35.0)),
        bituminous_thickness_mm=float(d.get("bituminous_thickness_mm", 100.0)),
        season_factor=float(d.get("season_factor", 1.0)),
        subgrade_type=d.get("subgrade_type", "granular"),
        design_traffic_msa=float(d.get("design_traffic_msa", 10.0)),
        road_category=d.get("road_category", "NH / SH"),
        notes=d.get("notes", "") or "",
    )
    return compute_overlay(inp)


def _rehydrate_cold_mix(row) -> "ColdMixResult | None":
    if not row or not row.inputs_json:
        return None
    try:
        d = json.loads(row.inputs_json)
    except json.JSONDecodeError:
        return None
    inp = ColdMixInput(
        aggregate_mass_kg=float(d.get("aggregate_mass_kg", 100.0)),
        emulsion_pct=float(d.get("emulsion_pct", 8.0)),
        emulsion_residue_pct=float(d.get("emulsion_residue_pct", 60.0)),
        water_addition_pct=float(d.get("water_addition_pct", 4.0)),
        filler_pct=float(d.get("filler_pct", 2.0)),
        mix_type=d.get("mix_type", "Dense-Graded"),
        notes=d.get("notes", "") or "",
    )
    return compute_cold_mix(inp)


def _rehydrate_traffic(row) -> "TrafficResult | None":
    if not row or not row.inputs_json:
        return None
    try:
        d = json.loads(row.inputs_json)
    except json.JSONDecodeError:
        return None
    inp = TrafficInput(
        initial_cvpd=float(d.get("initial_cvpd", 2000.0)),
        growth_rate_pct=float(d.get("growth_rate_pct", 7.5)),
        design_life_years=int(d.get("design_life_years", 15)),
        terrain=d.get("terrain", "Plain"),
        lane_config=d.get("lane_config", "Two-lane carriageway"),
        vdf=d.get("vdf"),
        ldf=d.get("ldf"),
        road_category=d.get("road_category", "NH / SH"),
        notes=d.get("notes", "") or "",
        axle_spectrum_kn=tuple(d.get("axle_spectrum_kn") or ()),
        wim_records=tuple(d.get("wim_records") or ()),
        survey_source_file=d.get("survey_source_file", "") or "",
    )
    return compute_traffic_analysis(inp)


def _rehydrate_condition(row) -> "ConditionSurveyResult | None":
    if not row or not row.inputs_json:
        return None
    try:
        d = json.loads(row.inputs_json)
    except json.JSONDecodeError:
        return None
    records = tuple(
        DistressRecord(
            distress_type=r.get("distress_type", "cracking"),
            severity=r.get("severity", "low"),
            length_m=float(r.get("length_m") or 0),
            area_m2=float(r.get("area_m2") or 0),
            count=int(r.get("count") or 0),
            notes=r.get("notes", "") or "",
            image_paths=tuple(r.get("image_paths") or ()),
        )
        for r in (d.get("records") or ())
    )
    inp = ConditionSurveyInput(
        work_name=d.get("work_name", "") or "",
        surveyed_by=d.get("surveyed_by", "") or "",
        survey_date=d.get("survey_date", "") or "",
        chainage_from_km=float(d.get("chainage_from_km") or 0),
        chainage_to_km=float(d.get("chainage_to_km") or 0),
        lane_id=d.get("lane_id", "") or "",
        records=records,
        notes=d.get("notes", "") or "",
        image_paths=tuple(d.get("image_paths") or ()),
        ai_classification_hint=d.get("ai_classification_hint", "") or "",
        gis_geometry_geojson=d.get("gis_geometry_geojson", "") or "",
    )
    return compute_condition_survey(inp)


def _rehydrate_material_qty(row) -> "MaterialQuantityResult | None":
    if not row or not row.inputs_json:
        return None
    try:
        d = json.loads(row.inputs_json)
    except json.JSONDecodeError:
        return None
    layers_raw = d.get("layers") or []
    layers = tuple(
        LayerInput(
            layer_type=L.get("layer_type", "DBM"),
            length_m=float(L.get("length_m", 1000.0)),
            width_m=float(L.get("width_m", 3.5)),
            thickness_mm=float(L.get("thickness_mm", 40.0)),
            density_t_m3=L.get("density_t_m3"),
            binder_pct=L.get("binder_pct"),
            spray_rate_kgm2=L.get("spray_rate_kgm2"),
            waste_pct=float(L.get("waste_pct", 2.0)),
            notes=L.get("notes", "") or "",
        )
        for L in layers_raw
    )
    return compute_material_quantity(MaterialQuantityInput(
        project_id=d.get("project_id"), layers=layers,
        notes=d.get("notes", "") or "",
    ))


def _rehydrate_micro(row) -> "MicroSurfacingResult | None":
    if not row or not row.inputs_json:
        return None
    try:
        d = json.loads(row.inputs_json)
    except json.JSONDecodeError:
        return None
    inp = MicroSurfacingInput(
        surfacing_type=d.get("surfacing_type", "Type II"),
        aggregate_mass_kg=float(d.get("aggregate_mass_kg", 100.0)),
        emulsion_pct=float(d.get("emulsion_pct", 13.0)),
        emulsion_residue_pct=float(d.get("emulsion_residue_pct", 62.0)),
        additive_water_pct=float(d.get("additive_water_pct", 8.0)),
        mineral_filler_pct=float(d.get("mineral_filler_pct", 1.5)),
        notes=d.get("notes", "") or "",
    )
    return compute_micro_surfacing(inp)


# ---------------------------------------------------------------------------
# Section dispatch (sub_module key -> writer)
# ---------------------------------------------------------------------------

def _maint_ctx(ctx: CombinedReportContext) -> MaintenanceReportContext:
    return MaintenanceReportContext(
        project_title=ctx.project_title,
        work_name=ctx.work_name,
        work_order_no=ctx.work_order_no,
        work_order_date=ctx.work_order_date,
        client=ctx.client,
        agency=ctx.agency,
        submitted_by=ctx.submitted_by,
        lab_name=ctx.lab_name,
        report_date=ctx.report_date,
    )


def _struct_ctx(ctx: CombinedReportContext) -> StructuralReportContext:
    return StructuralReportContext(
        project_title=ctx.project_title,
        work_name=ctx.work_name,
        work_order_no=ctx.work_order_no,
        work_order_date=ctx.work_order_date,
        client=ctx.client,
        agency=ctx.agency,
        submitted_by=ctx.submitted_by,
        lab_name=ctx.lab_name,
        report_date=ctx.report_date,
    )


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def build_combined_report(
    out_path: Path,
    db,
    project_id: int,
    ctx: CombinedReportContext,
    *,
    mix_result_live: "MixDesignResult | ImportedMixResult | None" = None,
    mix_chart_set: Optional[MarshallChartSet] = None,
    mix_material_calc=None,
) -> tuple[Path, list[str]]:
    """Build a single Word document for every module that has saved data.

    Returns ``(output_path, included_sections)``.
    The mix-design section is included only if a live ``MixDesignResult``
    is supplied — the DB summary alone does not have the full chart payload.
    """
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    p = db.get_project(project_id)
    if p is None:
        raise ValueError(f"Project #{project_id} not found.")

    # Discover persisted modules
    sd_row = db.latest_structural_design(project_id)
    ov_row = db.latest_maintenance_design(project_id, "overlay")
    cm_row = db.latest_maintenance_design(project_id, "cold_mix")
    ms_row = db.latest_maintenance_design(project_id, "micro_surfacing")

    structural = _rehydrate_structural(sd_row)
    overlay = _rehydrate_overlay(ov_row)
    cold_mix = _rehydrate_cold_mix(cm_row)
    micro = _rehydrate_micro(ms_row)
    mq_row = db.latest_material_quantity(project_id)
    material_qty = _rehydrate_material_qty(mq_row)
    tr_row = db.latest_traffic_analysis(project_id)
    traffic = _rehydrate_traffic(tr_row)
    cs_row = db.latest_condition_survey(project_id)
    condition = _rehydrate_condition(cs_row)

    have_mix = mix_result_live is not None
    have_any = any((have_mix, traffic, structural, overlay, cold_mix, micro,
                    material_qty, condition))
    if not have_any:
        raise ValueError(
            "No module data found for this project — compute and save at "
            "least one module before exporting a combined report."
        )

    doc = new_portrait_document()

    # ---- Title page ----
    add_heading(doc, "COMBINED PAVEMENT-DESIGN REPORT",
                level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    if ctx.project_title:
        add_p(doc, ctx.project_title, bold=True, size=12,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"{ctx.lab_name}  •  Report Date: {ctx.report_date}",
          size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "Project Information", level=2)
    add_kv_table(doc, (
        ("Name of Work",     ctx.work_name),
        ("Work Order No.",   ctx.work_order_no),
        ("Work Order Date",  ctx.work_order_date),
        ("Client",           ctx.client),
        ("Agency",           ctx.agency),
        ("Submitted By",     ctx.submitted_by),
        ("Binder Grade",     ctx.binder_grade),
    ))

    # Contents preview
    add_heading(doc, "Contents of this Report", level=2)
    toc_rows: list[list[str]] = []
    if have_mix:
        toc_rows.append(["Bituminous Mix Design",
                         "MoRTH Section 500 / IRC:111 / Marshall Mix Design"])
    if traffic:
        toc_rows.append(["Traffic / ESAL / MSA Analysis", "IRC:37-2018 / AASHTO-1993"])
    if structural:
        toc_rows.append(["Flexible Pavement Structural Design",
                         "IRC:37-2018"])
    if overlay:
        toc_rows.append(["Overlay Design (BBD)", "IRC:81-1997"])
    if cold_mix:
        toc_rows.append(["Cold Mix Design", "IRC:SP:100-2014"])
    if micro:
        toc_rows.append(["Micro-Surfacing Design", "IRC:SP:81"])
    if material_qty:
        toc_rows.append(["Bill of Material Quantities",
                         "MoRTH-400 / MoRTH-500 / IRC:111"])
    if condition:
        toc_rows.append(["Pavement Condition Survey",
                         "ASTM D6433 / IRC:82-1982 (placeholder PCI)"])
    # Phase 12 synthesis — derived on-demand from the rehydrated
    # condition + traffic + maintenance results. Not persisted (no DB
    # schema change in Phase 15 P2).
    rehab_synthesis = None
    if condition is not None:
        from app.core import (
            RecommendationContext,
            compute_rehab_recommendations,
        )
        rehab_synthesis = compute_rehab_recommendations(
            RecommendationContext(
                condition=condition,
                traffic=traffic,
                overlay_design=overlay,
                cold_mix_design=cold_mix,
                micro_surfacing_design=micro,
            )
        )
        toc_rows.append([
            "Rehabilitation Recommendations",
            "IRC:82-1982 / IRC:81-1997 / IRC:115 / IRC:SP:81 / IRC:SP:101",
        ])

    from ._docx_common import add_table
    add_table(doc, ["Section", "Governing Reference"], toc_rows)

    included: list[str] = []

    # ---- Mix design (uses existing word_report internals) ----
    if have_mix:
        from . import word_report
        from app.core import MIX_SPECS

        # Need a chart set if not supplied
        chart_set = mix_chart_set or build_chart_set(
            mix_result_live.summary, mix_result_live.obc)
        chart_image_dir = out_path.parent / f"{out_path.stem}_charts"

        doc.add_page_break()
        # Build a temporary mix-design doc to source content from? Simpler:
        # we directly call the existing helpers to draw the section onto our
        # current doc — but build_mix_design_docx is a one-shot file builder.
        # Cleanest path: build the mix-design docx separately and tell the
        # user the combined report will reference it. Better path: refactor
        # word_report to expose a section writer. We do the latter inline
        # via the existing public function on a sub-doc, then we won't
        # double-load python-docx images here — just emit a short bridge
        # section pointing to the standalone mix-design file generated next
        # to the combined doc.
        mix_path = out_path.with_name(f"{out_path.stem}_MixDesign.docx")
        # F4: refuse to render a mix-design report under the hidden DBM-II
        # default. If a live mix-design result is being included, the caller
        # must supply mix_type_key explicitly.
        if not ctx.mix_type_key:
            raise ValueError(
                "Combined report includes a live mix-design result but "
                "ctx.mix_type_key is empty. Set ReportContext.mix_type_key "
                "to the project's mix type (e.g. 'BC-II') before calling "
                "build_combined_report()."
            )
        word_report.build_mix_design_docx(
            mix_path,
            word_report.ReportContext(
                project_title=ctx.project_title,
                mix_type_key=ctx.mix_type_key,
                work_name=ctx.work_name,
                work_order_no=ctx.work_order_no,
                work_order_date=ctx.work_order_date,
                client=ctx.client,
                agency=ctx.agency,
                submitted_by=ctx.submitted_by,
                lab_name=ctx.lab_name,
                report_date=ctx.report_date,
                binder_grade=ctx.binder_grade,
                binder_properties=ctx.binder_properties,
                materials={},
            ),
            mix_result_live,
            chart_set,
            chart_image_dir,
            material_calc=mix_material_calc,
        )
        add_heading(doc, "Bituminous Mix Design", level=1,
                    align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(doc,
              "Mix design details — including 6 Marshall charts, gradation, "
              "specific-gravity tables, OBC and compliance — are issued as a "
              f"companion file: {mix_path.name}. Refer to that document for "
              "the full Bituminous Mix Design section. Design basis: MoRTH "
              "Section 500 / IRC:111.",
              size=10)
        # OBC summary inline so the combined doc is self-explanatory.
        # F4: mix_type_key is guaranteed non-empty here (raised above).
        obc = mix_result_live.obc
        spec_name = MIX_SPECS.get(ctx.mix_type_key)
        spec_label = spec_name.name if spec_name else ctx.mix_type_key
        add_kv_table(doc, (
            ("Mix Type", spec_label),
            ("Optimum Bitumen Content (OBC)", f"{obc.obc_pct:.2f} %"),
            ("Target Air Voids", f"{obc.target_air_voids_pct:.1f} %"),
            ("Bulk SG (Gsb)", f"{mix_result_live.bulk_sg_blend:.3f}"),
            ("Compliance",
             "PASS" if mix_result_live.compliance.overall_pass else "FAIL"),
        ))
        included.append("Bituminous Mix Design (companion file)")

    # ---- Traffic (precedes structural — its MSA feeds the structural design) ----
    if traffic:
        doc.add_page_break()
        write_traffic_section(doc, TrafficReportContext(
            project_title=ctx.project_title, work_name=ctx.work_name,
            work_order_no=ctx.work_order_no, work_order_date=ctx.work_order_date,
            client=ctx.client, agency=ctx.agency, submitted_by=ctx.submitted_by,
            lab_name=ctx.lab_name, report_date=ctx.report_date,
        ), traffic, include_header=True)
        included.append("Traffic / ESAL / MSA Analysis")

    # ---- Structural ----
    if structural:
        doc.add_page_break()
        write_structural_section(doc, _struct_ctx(ctx), structural,
                                 include_header=True)
        included.append("Flexible Pavement Structural Design")

    # ---- Maintenance sections ----
    maint_ctx = _maint_ctx(ctx)
    if overlay:
        doc.add_page_break()
        add_heading(doc, "MAINTENANCE / REHABILITATION",
                    level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        write_overlay_section(doc, maint_ctx, overlay)
        included.append("Overlay Design (BBD)")
    if cold_mix:
        if not overlay:
            doc.add_page_break()
            add_heading(doc, "MAINTENANCE / REHABILITATION",
                        level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        write_cold_mix_section(doc, maint_ctx, cold_mix)
        included.append("Cold Mix Design")
    if micro:
        if not (overlay or cold_mix):
            doc.add_page_break()
            add_heading(doc, "MAINTENANCE / REHABILITATION",
                        level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        write_micro_surfacing_section(doc, maint_ctx, micro)
        included.append("Micro-Surfacing Design")

    if material_qty:
        doc.add_page_break()
        mq_ctx = MaterialQuantityReportContext(
            project_title=ctx.project_title,
            work_name=ctx.work_name,
            work_order_no=ctx.work_order_no,
            work_order_date=ctx.work_order_date,
            client=ctx.client, agency=ctx.agency,
            submitted_by=ctx.submitted_by,
            lab_name=ctx.lab_name, report_date=ctx.report_date,
        )
        write_material_quantity_section(doc, mq_ctx, material_qty,
                                        include_header=True)
        included.append("Bill of Material Quantities")

    # ---- Pavement Condition Survey (Phase 10 — placeholder PCI) ----
    if condition:
        doc.add_page_break()
        cs_ctx = ConditionReportContext(
            project_title=ctx.project_title,
            work_name=ctx.work_name,
            work_order_no=ctx.work_order_no,
            work_order_date=ctx.work_order_date,
            client=ctx.client, agency=ctx.agency,
            submitted_by=ctx.submitted_by,
            lab_name=ctx.lab_name, report_date=ctx.report_date,
        )
        write_condition_section(doc, cs_ctx, condition, include_header=True)
        included.append("Pavement Condition Survey")

    # ---- Rehabilitation Recommendations (Phase 12 synthesis) ----
    if rehab_synthesis is not None:
        doc.add_page_break()
        rh_ctx = RehabReportContext(
            project_title=ctx.project_title,
            work_name=ctx.work_name,
            work_order_no=ctx.work_order_no,
            work_order_date=ctx.work_order_date,
            client=ctx.client, agency=ctx.agency,
            submitted_by=ctx.submitted_by,
            lab_name=ctx.lab_name, report_date=ctx.report_date,
        )
        write_rehab_section(doc, rh_ctx, rehab_synthesis, include_header=True)
        included.append("Rehabilitation Recommendations")

    add_signature_block(doc)
    doc.save(out_path)
    return out_path, included
