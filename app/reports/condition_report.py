"""Pavement Condition Survey — Word section + standalone docx (Phase 10).

Reuses Phase-6 ``_docx_common`` helpers and the engine's ``REFERENCES``.
Every report includes a PLACEHOLDER banner because the PCI weights are
uncalibrated foundation values (see app/core/condition_survey).
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import IMAGES_DIR
from app.core import ConditionSurveyResult, DISTRESS_TYPES
from app.core.condition_survey import REFERENCES as ENGINE_REFS

from ._docx_common import (
    add_code_references,
    add_heading,
    add_image_grid,
    add_kv_table,
    add_note,
    add_p,
    add_placeholder_banner,
    add_signature_block,
    add_table,
    new_portrait_document,
)


@dataclass(frozen=True, slots=True)
class ConditionReportContext:
    project_title: str = ""
    work_name: str = ""
    work_order_no: str = ""
    work_order_date: str = ""
    client: str = ""
    agency: str = ""
    submitted_by: str = ""
    lab_name: str = "Pavement Laboratory"
    report_date: str = field(default_factory=lambda: datetime.now().strftime("%d-%b-%Y"))


def write_condition_section(
    doc: Document,
    ctx: ConditionReportContext,
    result: ConditionSurveyResult,
    *,
    include_header: bool = True,
) -> None:
    if include_header:
        add_heading(doc, "PAVEMENT CONDITION SURVEY",
                    level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        if ctx.project_title:
            add_p(doc, ctx.project_title, bold=True, size=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(doc,
              "Basis: ASTM D6433 deduct-value method (shape) + IRC:82 "
              "maintenance treatment matrix. PCI score and rehab "
              "recommendations are foundation-level placeholders pending "
              "calibration.",
              italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(doc, f"{ctx.lab_name}  -  Report Date: {ctx.report_date}",
              size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_heading(doc, "Project Information", level=2)
        add_kv_table(doc, (
            ("Name of Work",    ctx.work_name),
            ("Work Order No.",  ctx.work_order_no),
            ("Client",          ctx.client),
            ("Agency",          ctx.agency),
            ("Submitted By",    ctx.submitted_by),
        ))

    # Placeholder banner (always shown in Phase 10)
    if result.is_placeholder:
        add_placeholder_banner(doc, "[PLACEHOLDER] " + (result.notes or
            "PCI weights and rehab recommendations are uncalibrated."))

    inp = result.inputs

    add_heading(doc, "Survey Metadata", level=2)
    add_kv_table(doc, (
        ("Surveyed By",         inp.surveyed_by or "-"),
        ("Survey Date",         inp.survey_date or "-"),
        ("Lane / Carriageway",  inp.lane_id or "-"),
        ("Chainage From (km)",  f"{inp.chainage_from_km:g}"),
        ("Chainage To (km)",    f"{inp.chainage_to_km:g}"),
        ("Records Captured",    str(len(inp.records))),
    ))

    # ---- Distress summary table ----------------------------------------
    add_heading(doc, "Distress Records", level=2)
    if result.breakdown:
        rows = []
        for b in result.breakdown:
            t = DISTRESS_TYPES.get(b.distress_type)
            label = t.label if t else b.distress_type
            unit = b.extent_unit.replace("_", " ")
            rows.append([
                label,
                b.severity.title(),
                f"{b.extent_value:g} {unit}" if b.extent_unit else "-",
                f"{b.deduct_value:.2f}",
                b.recommendation.treatment,
            ])
        add_table(doc, ["Distress", "Severity", "Extent",
                        "Deduct Value", "Rehab (placeholder)"], rows)
    else:
        add_p(doc, "No distress records were captured for this survey.",
              italic=True, size=10)

    # ---- PCI summary ---------------------------------------------------
    add_heading(doc, "PCI Score", level=2)
    add_kv_table(doc, (
        ("Total Deduct (Sum DV)",  f"{result.total_deduct:.2f}"),
        ("PCI Score (0-100)",      f"{result.pci_score:.2f}"),
        ("Condition Category",     result.condition_category),
        ("Calibration Status",     "PLACEHOLDER (uncalibrated)"
                                   if result.is_placeholder else "Calibrated"),
    ))

    # ---- Rehab recommendations summary (sourced from breakdown) --------
    if result.breakdown:
        add_heading(doc, "Rehab Recommendations", level=2)
        rehab_rows = []
        seen: set[tuple[str, str]] = set()
        for b in result.breakdown:
            key = (b.distress_type, b.severity)
            if key in seen:
                continue
            seen.add(key)
            t = DISTRESS_TYPES.get(b.distress_type)
            label = t.label if t else b.distress_type
            ref = b.recommendation.reference
            ref_label = f"{ref.code_id}" + (f" {ref.clause}" if ref.clause else "")
            rehab_rows.append([
                label, b.severity.title(),
                b.recommendation.treatment,
                ref_label,
            ])
        add_table(doc, ["Distress", "Severity", "Treatment (placeholder)",
                        "Reference"], rehab_rows)

    # ---- Image Evidence (Phase 11 step 3) ------------------------------
    add_heading(doc, "Image Evidence", level=2)
    total_embedded = 0
    if inp.image_paths:
        add_heading(doc, "Survey-wide", level=3)
        embedded = add_image_grid(doc, inp.image_paths,
                                  base_dir=IMAGES_DIR, cols=2, width_in=3.1)
        total_embedded += embedded
        if embedded == 0:
            add_p(doc, "(Referenced files not found on disk.)",
                  italic=True, size=10)
    # Per-distress sub-grids — iterate the original records so we keep
    # one section per captured record (the breakdown is per-row too).
    for idx, rec in enumerate(inp.records, start=1):
        if not rec.image_paths:
            continue
        t = DISTRESS_TYPES.get(rec.distress_type)
        label = t.label if t else rec.distress_type
        add_heading(doc,
                    f"Distress {idx} - {label} ({rec.severity.title()})",
                    level=3)
        embedded = add_image_grid(doc, rec.image_paths,
                                  base_dir=IMAGES_DIR, cols=2, width_in=2.4)
        total_embedded += embedded
        if embedded == 0:
            add_p(doc, "(Referenced files not found on disk.)",
                  italic=True, size=10)
    if total_embedded == 0 and not inp.image_paths and not any(
            r.image_paths for r in inp.records):
        add_p(doc, "No image evidence attached.", italic=True, size=10)

    # ---- Reserved-for-future placeholders ------------------------------
    add_heading(doc, "Reserved for Future Expansion", level=3)
    add_table(doc, ["Item", "Status"], [
        ["AI-assisted classification hint",
            inp.ai_classification_hint or "Not provided"],
        ["GIS geometry (GeoJSON)",
            "supplied" if inp.gis_geometry_geojson else "Not provided"],
    ])

    add_code_references(doc, ENGINE_REFS, heading="References")
    if result.notes and not result.is_placeholder:
        add_note(doc, result.notes)


def build_condition_docx(
    out_path: Path,
    ctx: ConditionReportContext,
    result: ConditionSurveyResult,
) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = new_portrait_document()
    write_condition_section(doc, ctx, result, include_header=True)
    add_signature_block(doc)
    doc.save(out_path)
    return out_path
