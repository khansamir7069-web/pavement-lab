"""Traffic Analysis — Word section + standalone docx (Phase 8).

Reuses Phase-6 ``_docx_common`` helpers and the engine's ``REFERENCES``.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core import TrafficResult
from app.core.traffic import REFERENCES as ENGINE_REFS

from ._docx_common import (
    add_code_references,
    add_heading,
    add_kv_table,
    add_note,
    add_p,
    add_signature_block,
    add_table,
    new_portrait_document,
)


@dataclass(frozen=True, slots=True)
class TrafficReportContext:
    project_title: str = ""
    work_name: str = ""
    work_order_no: str = ""
    work_order_date: str = ""
    client: str = ""
    agency: str = ""
    submitted_by: str = ""
    lab_name: str = "Pavement Laboratory"
    report_date: str = field(default_factory=lambda: datetime.now().strftime("%d-%b-%Y"))


def write_traffic_section(
    doc: Document, ctx: TrafficReportContext, result: TrafficResult,
    *, include_header: bool = True,
) -> None:
    if include_header:
        add_heading(doc, "TRAFFIC / ESAL / MSA ANALYSIS",
                    level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        if ctx.project_title:
            add_p(doc, ctx.project_title, bold=True, size=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(doc,
              "Design basis: IRC:37-2018 (cl. 4.6 MSA, Table 1 VDF, cl. 4.4 LDF). "
              "AASHTO ESAL reported for cross-reference.",
              italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p(doc, f"{ctx.lab_name}  •  Report Date: {ctx.report_date}",
              size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_heading(doc, "Project Information", level=2)
        add_kv_table(doc, (
            ("Name of Work",    ctx.work_name),
            ("Work Order No.",  ctx.work_order_no),
            ("Client",          ctx.client),
            ("Agency",          ctx.agency),
            ("Submitted By",    ctx.submitted_by),
        ))

    inp = result.inputs

    add_heading(doc, "Inputs", level=2)
    add_kv_table(doc, (
        ("Road Category",                    inp.road_category),
        ("Terrain",                          inp.terrain),
        ("Lane Configuration",               inp.lane_config),
        ("Initial Commercial Traffic (A)",   f"{inp.initial_cvpd:g} CVPD"),
        ("Annual Growth Rate (r)",           f"{inp.growth_rate_pct:g} %"),
        ("Design Life (n)",                  f"{inp.design_life_years} yr"),
        ("VDF (F) — user / preset",
            f"{inp.vdf:g}" if inp.vdf is not None
            else f"{result.vdf_used:g} (IRC:37 Table 1 preset)"),
        ("LDF (D) — user / preset",
            f"{inp.ldf:g}" if inp.ldf is not None
            else f"{result.ldf_used:g} (IRC:37 cl. 4.4 preset)"),
    ))

    add_heading(doc, "Computed Quantities", level=2)
    add_p(doc, "N = (365 · A · [(1+r)^n − 1]/r) · D · F · 10⁻⁶   (IRC:37 cl. 4.6)",
          italic=True, size=11)
    add_table(doc, ["Quantity", "Symbol", "Value", "Unit"], [
        ["Growth Factor",                "GF", f"{result.growth_factor:.4f}", "—"],
        ["VDF used",                     "F",  f"{result.vdf_used:g}",        "—"],
        ["LDF used",                     "D",  f"{result.ldf_used:g}",        "—"],
        ["Cumulative Design Traffic",    "N",  f"{result.design_msa:.2f}",    "MSA"],
        ["AASHTO ESAL (alias N × 10⁶)",  "—",  f"{result.aashto_esal:,.0f}",  "ESAL"],
        ["Traffic Category",             "—",  result.traffic_category,       "—"],
    ])

    # Reserved-for-future placeholders shown as a status table
    add_heading(doc, "Reserved for Future Expansion", level=3)
    add_table(doc, ["Item", "Status"], [
        ["Axle-Load Spectrum",     "Not provided" if not inp.axle_spectrum_kn else "supplied"],
        ["Weigh-In-Motion Records","Not provided" if not inp.wim_records else "supplied"],
        ["Traffic-Survey Source",  inp.survey_source_file or "Not provided"],
    ])

    add_code_references(doc, ENGINE_REFS, heading="References")
    if result.notes:
        add_note(doc, result.notes)


def build_traffic_docx(
    out_path: Path, ctx: TrafficReportContext, result: TrafficResult,
) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = new_portrait_document()
    write_traffic_section(doc, ctx, result, include_header=True)
    add_signature_block(doc)
    doc.save(out_path)
    return out_path
