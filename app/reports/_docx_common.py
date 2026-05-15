"""Shared python-docx helpers for Phase 6 module-aware report builders.

Mirrors the style choices in ``word_report.py`` (the existing mix-design
report) so all module sections look consistent. Re-exports are kept narrow
to avoid drift if ``word_report.py`` is later updated.
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


HEADING_FONT = "Arial"
BODY_FONT = "Arial"
ACCENT_RGB = RGBColor(0x1F, 0x3A, 0x68)
ACCENT_HEX = "1F3A68"
MUTED_HEX = "F4F6FA"


def configure_section(section) -> None:
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)


def add_page_number_footer(section) -> None:
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("Page ")
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)


def add_heading(doc: Document, text: str, level: int = 1,
                align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = HEADING_FONT
    run.font.bold = True
    run.font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
    if level <= 2:
        run.font.color.rgb = ACCENT_RGB


def add_p(doc: Document, text: str, bold: bool = False, size: int = 11,
          align=WD_ALIGN_PARAGRAPH.LEFT, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _style_cell_text(cell, *, bold: bool = False, size: int = 10,
                     color: str | None = None,
                     align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in cell.paragraphs:
        p.alignment = align
        for r in p.runs:
            r.font.name = BODY_FONT
            r.font.size = Pt(size)
            r.font.bold = bold
            if color:
                r.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              first_col_align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        _set_cell_shading(c, ACCENT_HEX)
        _style_cell_text(c, bold=True, size=10, color="FFFFFF")
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            c = t.rows[r_idx].cells[c_idx]
            c.text = str(val)
            align = first_col_align if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _style_cell_text(c, bold=False, size=10, align=align)


def add_kv_table(doc: Document, rows: Iterable[tuple[str, str]]) -> None:
    """Two-column 'Item / Detail' table for project metadata blocks."""
    rows_l = [[k, v] for k, v in rows if v]
    if not rows_l:
        return
    add_table(doc, ["Item", "Detail"], rows_l)


def add_note(doc: Document, text: str) -> None:
    """Italic muted paragraph used for code-citation footnotes."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x5B, 0x6E)


def add_placeholder_banner(doc: Document, text: str) -> None:
    """Bold orange-toned warning paragraph for placeholder/un-IRC-verified
    content. Single source of truth — every report uses this helper.
    """
    if not text:
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x8A, 0x5A, 0x00)


def add_code_references(doc: Document, refs, heading: str = "References") -> None:
    """Uniform citation block rendered from a tuple of CodeRef objects.

    ``refs`` is any iterable of ``app.core.CodeRef`` (or compatible objects
    exposing ``.label()``). Empty / falsy input is a no-op.
    """
    items = list(refs or ())
    if not items:
        return
    add_heading(doc, heading, level=3)
    rows = [[r.code_id, (r.clause or "—"),
             (getattr(get_code_record(r.code_id), "title", "") or r.note or "")]
            for r in items]
    add_table(doc, ["Code", "Clause / Table", "Title / Note"], rows)


def get_code_record(code_id: str):
    """Lazy import to avoid circular dep at module-load time."""
    from app.core.code_refs import get_code
    return get_code(code_id)


def add_image_grid(
    doc: Document,
    image_paths: Iterable[str],
    *,
    base_dir: Path,
    cols: int = 2,
    width_in: float = 3.1,
) -> int:
    """Render IMAGES_DIR-relative paths as a centered N-column python-docx
    table of pictures. Missing files are skipped silently (the panel
    already flags them).

    Returns the number of images actually embedded so callers can decide
    whether to render a heading or an empty-state note.
    """
    resolved: list[Path] = []
    for rel in image_paths or ():
        if not rel:
            continue
        p = base_dir / Path(rel)
        if p.is_file():
            resolved.append(p)
    if not resolved:
        return 0
    cols = max(1, int(cols))
    rows = (len(resolved) + cols - 1) // cols
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, p in enumerate(resolved):
        r, c = divmod(idx, cols)
        cell = t.rows[r].cells[c]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        try:
            run.add_picture(str(p), width=Inches(width_in))
        except Exception:
            # Corrupt / unreadable file — silently skip cell content.
            continue
    return len(resolved)


def add_signature_block(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_table(rows=1, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig.rows[0].cells[0].text = "Prepared By\n\n______________________"
    sig.rows[0].cells[1].text = "Approved By\n\n______________________"
    for cell in sig.rows[0].cells:
        _style_cell_text(cell, size=10)


def new_portrait_document() -> Document:
    """Document with consistent A4 portrait sections and page-number footer."""
    doc = Document()
    for section in doc.sections:
        configure_section(section)
        add_page_number_footer(section)
    return doc
