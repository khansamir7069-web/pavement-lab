"""Word report generator matching the consultancy sample layout.

Sections produced:
  1. Title block (mix-type heading + project info)
  2. Materials table (name | source)
  3. Aggregate gradation per fraction (% passing × sieve)
  4. Blend proportions
  5. Combined gradation vs MoRTH spec
  6. Specific gravities of materials
  7. Marshall mix design summary (5 rows)
  8. OBC + properties at OBC + compliance
  9. 6 Marshall design curves (2 × 3 grid)
 10. Signature block
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Mapping

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from app.core import MIX_SPECS, MIX_TYPES, MaterialCalcResult, MixDesignResult
from app.core.import_summary import ImportedMixResult
from app.graphs import MarshallChartSet, save_chart_pngs

HEADING_FONT = "Arial"
BODY_FONT = "Arial"
ACCENT_RGB = RGBColor(0x1F, 0x3A, 0x68)
ACCENT_HEX = "1F3A68"


@dataclass(frozen=True, slots=True)
class ReportContext:
    project_title: str
    mix_type_key: str
    work_order_no: str = ""
    work_order_date: str = ""
    work_name: str = ""
    client: str = ""
    agency: str = ""
    submitted_by: str = ""
    lab_name: str = "Pavement Laboratory"
    materials: Mapping[str, str] = field(default_factory=dict)
    report_date: str = field(default_factory=lambda: datetime.now().strftime("%d-%b-%Y"))
    binder_grade: str = ""                                          # e.g. "VG-30"
    binder_properties: Mapping[str, float] = field(default_factory=dict)


# ---------- docx helpers ----------------------------------------------------

def _add_heading(doc: Document, text: str, level: int = 1,
                 align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = HEADING_FONT
    run.font.bold = True
    run.font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
    if level <= 2:
        run.font.color.rgb = ACCENT_RGB


def _add_p(doc: Document, text: str, bold: bool = False, size: int = 11,
           align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    r.font.bold = bold


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _style_cell_text(cell, *, bold: bool = False, size: int = 10,
                     color: str | None = None,
                     align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in cell.paragraphs:
        p.alignment = align
        for r in p.runs:
            r.font.name = BODY_FONT
            r.font.size = Pt(size)
            r.font.bold = bold
            if color:
                r.font.color.rgb = RGBColor.from_string(color)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]],
               first_col_align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        _set_cell_shading(c, ACCENT_HEX)
        _style_cell_text(c, bold=True, size=10, color="FFFFFF")
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            c = t.rows[r_idx].cells[c_idx]
            c.text = str(val)
            align = first_col_align if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _style_cell_text(c, bold=False, size=10, align=align)


def _add_image_grid(doc: Document, image_paths: list[Path], cols: int = 2,
                    width_in: float = 3.1) -> None:
    rows = (len(image_paths) + cols - 1) // cols
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, p in enumerate(image_paths):
        r, c = divmod(idx, cols)
        cell = t.rows[r].cells[c]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(p), width=Inches(width_in))


def _add_page_number_footer(section) -> None:
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("Page ")
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)


# ---------- public api ------------------------------------------------------

def build_mix_design_docx(
    out_path: Path,
    ctx: ReportContext,
    result: "MixDesignResult | ImportedMixResult",
    chart_set: MarshallChartSet,
    chart_image_dir: Path | None = None,
    material_calc: MaterialCalcResult | None = None,
) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    chart_image_dir = chart_image_dir or out_path.parent / f"{out_path.stem}_charts"
    chart_paths = save_chart_pngs(chart_set, chart_image_dir)

    doc = Document()

    for section in doc.sections:
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        _add_page_number_footer(section)

    # ----- Title block
    spec = MIX_SPECS[ctx.mix_type_key]
    _add_heading(doc, f"REPORT ON MIX DESIGN OF {spec.name.upper()}",
                 level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    if ctx.project_title:
        _add_p(doc, ctx.project_title, bold=True, size=12,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_p(doc, f"{ctx.lab_name}  •  Report Date: {ctx.report_date}",
           size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # F2 placeholder warning — surfaces unverified spec status in the report.
    _record = MIX_TYPES.get(ctx.mix_type_key)
    if _record and (_record.status or "").strip() == "placeholder_editable":
        from ._docx_common import add_placeholder_banner
        _src = _record.applicable_code or "unverified"
        add_placeholder_banner(
            doc,
            f"⚠ Spec limits for {_record.mix_code} ({_record.full_name}) are "
            f"not IRC-verified. Compliance verdict in this report is "
            f"indicative only — confirm against the relevant IRC clause "
            f"before adoption. (Source: {_src})"
        )

    # Project info
    _add_heading(doc, "Project Information", level=2)
    proj_rows: list[list[str]] = []
    if ctx.work_name: proj_rows.append(["Name of Work", ctx.work_name])
    if ctx.work_order_no: proj_rows.append(["Work Order No.", ctx.work_order_no])
    if ctx.work_order_date: proj_rows.append(["Work Order Date", ctx.work_order_date])
    if ctx.client: proj_rows.append(["Client", ctx.client])
    if ctx.agency: proj_rows.append(["Agency", ctx.agency])
    if ctx.submitted_by: proj_rows.append(["Submitted By", ctx.submitted_by])
    if ctx.binder_grade: proj_rows.append(["Binder Grade", ctx.binder_grade])
    # F3 (Phase-9 audit close-out): show the compaction blows per face
    # for the selected mix so the engineer / lab can match specimen prep.
    proj_rows.append([
        "Compaction (each face)",
        f"{spec.compaction_blows_each_face} blows ({_record.applicable_code or '—'})"
        if _record else f"{spec.compaction_blows_each_face} blows",
    ])
    if proj_rows:
        _add_table(doc, ["Item", "Detail"], proj_rows)

    # Binder properties table (only if any provided)
    if ctx.binder_properties:
        from app.core import PROPERTY_LABELS
        _add_heading(doc, "Binder Properties", level=3)
        rows = []
        for key, val in ctx.binder_properties.items():
            if key.startswith("_"):
                continue
            label = PROPERTY_LABELS.get(key, key.replace("_", " ").title())
            rows.append([label, f"{val:g}" if isinstance(val, (int, float)) else str(val)])
        if rows:
            _add_table(doc, ["Property", "Value"], rows)
        notes = ctx.binder_properties.get("_notes", "")
        if notes:
            _add_p(doc, f"Notes: {notes}", size=10)

    # ----- 1. Materials
    _add_heading(doc, "1. Materials", level=2)
    if ctx.materials:
        rows = [[name, src or "—"] for name, src in ctx.materials.items()]
        _add_table(doc, ["Material", "Source / Brand"], rows)

    # Detect whether this result has raw gradation / SG data
    _has_raw = getattr(result, "gradation", None) is not None

    # ----- 2. Aggregate gradation per fraction
    if _has_raw:
        _add_heading(doc, "2. Aggregate Gradation (% Passing)", level=2)
        grad = result.gradation
        agg_names = list(grad.pass_pct.keys())
        headers = ["IS Sieve (mm)"] + agg_names
        rows = []
        for i, sieve in enumerate(grad.sieve_sizes_mm):
            row = [f"{sieve:g}"]
            for name in agg_names:
                row.append(f"{grad.pass_pct[name][i]:.1f}")
            rows.append(row)
        _add_table(doc, headers, rows)

        # ----- 3. Blend proportions
        _add_heading(doc, "3. Blend Proportions", level=2)
        blend_rows = [[name, f"{ratio * 100:.1f}%"]
                      for name, ratio in grad.blend_ratios.items()]
        blend_rows.append(["TOTAL", f"{grad.blend_ratio_sum * 100:.1f}%"])
        _add_table(doc, ["Aggregate", "Proportion"], blend_rows)
        _add_p(doc,
               f"Aggregates are mixed in the proportions above to meet the combined "
               f"gradation specified for {spec.name} (MoRTH).",
               size=10)

        # ----- 4. Combined gradation vs spec
        _add_heading(doc, "4. Combined Gradation vs Specification", level=2)
        rows = []
        for i, sieve in enumerate(grad.sieve_sizes_mm):
            rows.append([
                f"{sieve:g}",
                f"{grad.combined_pass_pct[i]:.2f}",
                f"{grad.spec_lower[i]:g} – {grad.spec_upper[i]:g}",
                "Within" if grad.within_spec[i] else "Out",
            ])
        _add_table(doc, ["IS Sieve (mm)", "Combined % Passing",
                         "MoRTH Spec", "Status"], rows)

        # ----- 5. Specific gravities
        _add_heading(doc, "5. Specific Gravities of Materials", level=2)
        sg_rows: list[list[str]] = []
        for name, sg in result.sg_coarse.items():
            sg_rows.append([f"Coarse {name}",
                           f"{sg.avg_bulk_ovendry:.3f}",
                           f"{sg.avg_absorption_pct:.2f}%"])
        for name, sg in result.sg_fine.items():
            sg_rows.append([f"Fine {name}", f"{sg.avg_bulk_ovendry:.3f}", "—"])
        sg_rows.append([
            f"Bitumen ({ctx.binder_grade or 'VG-30'})",
            f"{result.bitumen_sg:.3f}", "—"
        ])
        sg_rows.append(["Combined aggregate (Gsb)", f"{result.bulk_sg_blend:.3f}", "—"])
        _add_table(doc, ["Material", "Bulk Sp. Gr.", "Water Absorption"], sg_rows)
    else:
        # Imported from summary — raw section not available
        _add_heading(doc, "2–5. Gradation, Blend & Specific Gravities", level=2)
        _add_p(doc,
               "This report was generated from an imported summary table.  "
               "Raw gradation, blend proportions and specific-gravity test data "
               "are not available in this file.",
               size=10)
        _add_p(doc, f"Estimated Gsb (from VMA back-calc): {result.bulk_sg_blend:.4f}",
               size=10)

    # ----- 6. Marshall summary
    _add_heading(doc, "6. Marshall Mix Design Summary", level=2)
    rows = []
    for r in result.summary.rows:
        rows.append([
            f"{r.bitumen_pct:.1f}",
            f"{r.gmm:.3f}",
            f"{r.gmb:.3f}",
            f"{r.air_voids_pct:.2f}",
            f"{r.vma_pct:.2f}",
            f"{r.vfb_pct:.2f}",
            f"{r.stability_kn:.2f}",
            f"{r.flow_mm:.2f}",
            f"{r.marshall_quotient:.2f}",
        ])
    _add_table(doc,
               ["Pb %", "Gmm", "Gmb", "VIM %", "VMA %", "VFB %",
                "Stab. kN", "Flow mm", "MQ"],
               rows)

    # ----- 7. OBC and compliance
    _add_heading(doc, "7. Optimum Bitumen Content (OBC)", level=2)
    obc = result.obc
    # F5 (Phase-9 audit close-out): the target air-voids value below is
    # the midpoint of MIX_SPECS[mix].air_voids_min/max, not a flat 4.0
    # constant. Annotate the source so the reader can audit the choice.
    av_source = (
        f"midpoint of {spec.air_voids_min_pct:g}–{spec.air_voids_max_pct:g}% "
        f"from {_record.applicable_code or spec.name}"
        if _record else "default"
    )
    _add_p(doc,
           f"OBC = {obc.obc_pct:.2f}% (at target {obc.target_air_voids_pct:.1f}% "
           f"air voids — {av_source}; method: {obc.method.replace('_', ' ')})",
           bold=True, size=12)

    _add_heading(doc, "Mix Parameters at OBC vs Specification", level=3)
    rows = []
    for item in result.compliance.items:
        rows.append([item.name, f"{item.value:.2f}", item.requirement,
                     "PASS" if item.pass_ else "FAIL"])
    rows.append(["Gmb at OBC", f"{obc.gmb_at_obc:.3f}", "—", "—"])
    rows.append(["Gmm at OBC", f"{obc.gmm_at_obc:.3f}", "—", "—"])
    _add_table(doc, ["Property", "Value at OBC", "Requirement", "Status"], rows)

    overall = "PASS" if result.compliance.overall_pass else "FAIL"
    _add_p(doc, f"Overall compliance ({result.compliance.spec_name}): {overall}",
           bold=True, size=12)

    # ----- 8. Material Calculation
    if material_calc is not None:
        _add_heading(doc, "8. Material Calculation for Sample Preparation", level=2)
        _add_p(doc,
               "Quantities to weigh out when preparing a Marshall specimen of the "
               "standard sample (left bitumen content) and the target preparation "
               "sample (right bitumen content), with dry-material breakdown by "
               "aggregate fraction.",
               size=10)
        _add_table(
            doc,
            ["Block", "Pb %", "Aggregate %", "Aggregate wt (g)",
             "Bitumen wt (g)", "Total mix wt (g)"],
            [
                ["Standard sample",
                 f"{material_calc.standard.bitumen_pct:.2f}",
                 f"{material_calc.standard.aggregate_pct:.2f}",
                 f"{material_calc.standard.aggregate_weight_g:.2f}",
                 f"{material_calc.standard.bitumen_weight_g:.2f}",
                 f"{material_calc.standard.total_mix_weight_g:.2f}"],
                ["Target sample",
                 f"{material_calc.target.bitumen_pct:.2f}",
                 f"{material_calc.target.aggregate_pct:.2f}",
                 f"{material_calc.target.aggregate_weight_g:.2f}",
                 f"{material_calc.target.bitumen_weight_g:.2f}",
                 f"{material_calc.target.total_mix_weight_g:.2f}"],
            ],
        )
        _add_heading(doc, "Dry Material — Standard Sample", level=3)
        rows_std = []
        for row in material_calc.dry_material_standard:
            rows_std.append([row.name, f"{row.fraction * 100:.1f}%", f"{row.weight_g:.2f}"])
        rows_std.append(["TOTAL", "100.0%", f"{material_calc.total_dry_standard_g:.2f}"])
        _add_table(doc, ["Particle Size", "Blend Ratio", "Weight (g)"], rows_std)

        _add_heading(doc, "Dry Material — Target Sample", level=3)
        rows_tgt = []
        for row in material_calc.dry_material_target:
            rows_tgt.append([row.name, f"{row.fraction * 100:.1f}%", f"{row.weight_g:.2f}"])
        rows_tgt.append(["TOTAL", "100.0%", f"{material_calc.total_dry_target_g:.2f}"])
        _add_table(doc, ["Particle Size", "Blend Ratio", "Weight (g)"], rows_tgt)

    # ----- 9. Charts
    doc.add_page_break()
    _add_heading(doc, "9. Marshall Design Curves", level=2)
    image_paths = [chart_paths[cd.key] for cd in chart_set.charts]
    _add_image_grid(doc, image_paths, cols=2, width_in=3.1)

    # ----- 9. Signature
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_table(rows=1, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig.rows[0].cells[0].text = "Prepared By\n\n______________________"
    sig.rows[0].cells[1].text = "Approved By\n\n______________________"
    for cell in sig.rows[0].cells:
        _style_cell_text(cell, size=10)

    doc.save(out_path)
    return out_path


def export_to_pdf(docx_path: Path) -> Path:
    """Convert a .docx to .pdf using Word (docx2pdf) or LibreOffice fallback."""
    docx_path = Path(docx_path)
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return pdf_path
    except Exception:
        import shutil
        import subprocess

        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", str(docx_path.parent), str(docx_path)],
                check=True,
            )
            return pdf_path
        raise RuntimeError(
            "No PDF converter available. Install Microsoft Word or LibreOffice."
        )
